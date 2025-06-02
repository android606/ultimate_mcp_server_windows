# ultimate_mcp_server/tools/document_conversion_and_processing.py
"""Standalone Document Processing Toolkit functions for MCP Server.

A comprehensive, fault-tolerant toolkit for document processing, providing:
(Functionality remains the same as the original class docstring)
"""

from __future__ import annotations

###############################################################################
# Imports                                                                     #
###############################################################################
# Standard library imports
import asyncio
import base64
import csv
import functools
import hashlib
import html
import io
import json
import math
import os
import re
import tempfile
import textwrap
import time
from contextlib import contextmanager
from io import StringIO
from pathlib import Path
from typing import (
    TYPE_CHECKING,
    Any,
    Awaitable,
    Callable,
    Dict,
    List,
    Optional,
    Pattern,
    Sequence,
    Set,
    Tuple,
    Union,
)

# Third-party imports
# import html2text  # REMOVED - Using lazy import
# from bs4 import BeautifulSoup, Tag  # REMOVED - Using lazy import
# from rapidfuzz import fuzz  # REMOVED - Using lazy import

# Lazy import functions
def _get_bs4():
    """Lazy import for BeautifulSoup to avoid startup dependency."""
    try:
        from bs4 import BeautifulSoup, Tag
        return BeautifulSoup, Tag
    except ImportError as e:
        raise ImportError("BeautifulSoup4 package is not installed. Please install with: pip install beautifulsoup4")

def _get_html2text():
    """Lazy import for html2text to avoid startup dependency."""
    try:
        import html2text
        return html2text
    except ImportError as e:
        raise ImportError("html2text package is not installed. Please install with: pip install html2text")

def _get_rapidfuzz():
    """Lazy import for rapidfuzz to avoid startup dependency."""
    try:
        from rapidfuzz import fuzz
        return fuzz
    except ImportError as e:
        raise ImportError("rapidfuzz package is not installed. Please install with: pip install rapidfuzz")

# Local application imports
from ultimate_mcp_server.constants import Provider
from ultimate_mcp_server.exceptions import ProviderError, ToolError, ToolInputError
from ultimate_mcp_server.tools.base import (
    with_error_handling,
    with_retry,
    with_tool_metrics,
)
from ultimate_mcp_server.tools.completion import generate_completion
from ultimate_mcp_server.utils import get_logger

# Type checking imports
if TYPE_CHECKING:
    import numpy as np
    import pandas as pd
    import tiktoken
    from docling.datamodel.pipeline_options import AcceleratorDevice as _AcceleratorDeviceType
    from docling_core.types.doc import DoclingDocument as _DoclingDocumentType
    from docling_core.types.doc import ImageRefMode as _ImageRefModeType
    from PIL import Image as PILImage
    from tiktoken import Encoding

# ───────────────────── Optional Dependency Check & Initialization ───────────────────
_DOCLING_AVAILABLE = False
try:
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import (
        AcceleratorDevice,
        AcceleratorOptions,
        PdfPipelineOptions,
    )
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling_core.types.doc import DoclingDocument, ImageRefMode

    _DOCLING_AVAILABLE = True
    _DoclingDocumentType = DoclingDocument
    _ImageRefModeType = ImageRefMode
    _AcceleratorDeviceType = AcceleratorDevice

except ImportError:
    _DoclingDocumentType = Any
    _ImageRefModeType = Any
    _AcceleratorDeviceType = Any
    InputFormat = None
    AcceleratorDevice = None
    AcceleratorOptions = None
    PdfPipelineOptions = None
    DocumentConverter = None
    PdfFormatOption = None
    pass

_PANDAS_AVAILABLE = False
try:
    import pandas as pd

    _PANDAS_AVAILABLE = True
except ModuleNotFoundError:
    pd = None

_TIKTOKEN_AVAILABLE = False
try:
    import tiktoken

    _TIKTOKEN_AVAILABLE = True
except ModuleNotFoundError:
    tiktoken = None

_PYPDF2_AVAILABLE = False
try:
    import PyPDF2

    _PYPDF2_AVAILABLE = True
except ImportError:
    PyPDF2 = None

_DOCX_AVAILABLE = False
try:
    import docx

    _DOCX_AVAILABLE = True
except ImportError:
    docx = None

_NUMPY_AVAILABLE = False
try:
    import numpy as np

    _NUMPY_AVAILABLE = True
except ImportError:
    np = None

_PIL_AVAILABLE = False
try:
    from PIL import Image, ImageEnhance, ImageFilter

    _PIL_AVAILABLE = True
except ImportError:
    Image, ImageEnhance, ImageFilter = None, None, None

_CV2_AVAILABLE = False
try:
    import cv2

    _CV2_AVAILABLE = True
except ImportError:
    cv2 = None

_PYTESSERACT_AVAILABLE = False
try:
    import pytesseract

    _PYTESSERACT_AVAILABLE = True
except ImportError:
    pytesseract = None

_PDF2IMAGE_AVAILABLE = False
try:
    from pdf2image import convert_from_bytes, convert_from_path

    _PDF2IMAGE_AVAILABLE = True
except ImportError:
    convert_from_bytes, convert_from_path = None, None

_PDFPLUMBER_AVAILABLE = False
try:
    import pdfplumber

    _PDFPLUMBER_AVAILABLE = True
except ImportError:
    pdfplumber = None

_PYMUPDF_AVAILABLE = False
try:
    import pymupdf  # PyMuPDF

    _PYMUPDF_AVAILABLE = True
except ImportError:
    pymupdf = None

_TRAFILATURA_AVAILABLE = False
try:
    import trafilatura

    _TRAFILATURA_AVAILABLE = True
except ImportError:
    trafilatura = None

_READABILITY_AVAILABLE = False
try:
    import readability

    _READABILITY_AVAILABLE = True
except ImportError:
    readability = None

try:
    from markdownify import markdownify as _markdownify_fallback
except ModuleNotFoundError:
    _markdownify_fallback = None

# ───────────────────── Module Level Logger ─────────────────────────
logger = get_logger("ultimate_mcp_server.tools.document_processing")

# ───────────────────── Module Level Config & State ──────────────────
DEFAULT_EXTRACTION_STRATEGY = "hybrid_direct_ocr"
_VALID_FORMATS = {"markdown", "text", "html", "json", "doctags"}
_OCR_COMPATIBLE_FORMATS = {"text", "markdown"}
_VALID_EXTRACTION_STRATEGIES = {
    "docling",
    "direct_text",
    "ocr",
    "hybrid_direct_ocr",
}

# Acceleration Device Mapping (Docling)
if _DOCLING_AVAILABLE and AcceleratorDevice:
    _ACCEL_MAP = {
        "auto": AcceleratorDevice.AUTO,
        "cpu": AcceleratorDevice.CPU,
        "cuda": AcceleratorDevice.CUDA,
        "mps": AcceleratorDevice.MPS,
    }
else:
    _ACCEL_MAP = {"auto": "auto", "cpu": "cpu", "cuda": "cuda", "mps": "mps"}

# HTML Detection Patterns
_RE_FLAGS = re.MULTILINE | re.IGNORECASE
_HTML_PATTERNS: Sequence[Pattern] = [
    re.compile(p, _RE_FLAGS)
    for p in (
        r"<\s*[a-zA-Z]+[^>]*>",
        r"<\s*/\s*[a-zA-Z]+\s*>",
        r"&[a-zA-Z]+;",
        r"&#[0-9]+;",
        r"<!\s*DOCTYPE",
        r"<!\s*--",
    )
]

# Content Type Patterns (Used by detect_content_type)
_CONTENT_PATTERNS: Dict[str, List[Tuple[Pattern, float]]] = {
    "html": [
        (re.compile(r"<html", re.I), 5.0),
        (re.compile(r"<head", re.I), 4.0),
        (re.compile(r"<body", re.I), 4.0),
        (re.compile(r"</(div|p|span|a|li)>", re.I), 1.0),
        (re.compile(r"<[a-z][a-z0-9]*\s+[^>]*>", re.I), 0.8),
        (re.compile(r"<!DOCTYPE", re.I), 5.0),
        (re.compile(r"&\w+;"), 0.5),
    ],
    "markdown": [
        (re.compile(r"^#{1,6}\s+", re.M), 4.0),
        (re.compile(r"^\s*[-*+]\s+", re.M), 2.0),
        (re.compile(r"^\s*\d+\.\s+", re.M), 2.0),
        (re.compile(r"`[^`]+`"), 1.5),
        (re.compile(r"^```", re.M), 5.0),
        (re.compile(r"\*{1,2}[^*\s]+?\*{1,2}"), 1.0),
        (re.compile(r"!\[.*?\]\(.*?\)", re.M), 3.0),
        (re.compile(r"\[.*?\]\(.*?\)", re.M), 2.5),
        (re.compile(r"^>.*", re.M), 2.0),
        (re.compile(r"^-{3,}$", re.M), 3.0),
    ],
    "code": [
        (re.compile(r"def\s+\w+\(.*\):"), 3.0),
        (re.compile(r"class\s+\w+"), 3.0),
        (re.compile(r"import\s+|from\s+"), 3.0),
        (
            re.compile(r"((function\s+\w+\(|const|let|var)\s*.*?=>|\b(document|console|window)\.)"),
            3.0,
        ),
        (re.compile(r"public\s+|private\s+|static\s+"), 2.5),
        (re.compile(r"#include"), 3.0),
        (re.compile(r"<\?php"), 4.0),
        (re.compile(r"console\.log"), 2.0),
        (re.compile(r";\s*$"), 1.0),
        (re.compile(r"\b(var|let|const|int|float|string|bool)\b"), 1.5),
        (re.compile(r"//.*$"), 1.0),
        (re.compile(r"/\*.*?\*/", re.S), 1.5),
    ],
}
_LANG_PATTERNS: List[Tuple[Pattern, str]] = [
    (re.compile(r"(def\s+\w+\(.*?\):|import\s+|from\s+\S+\s+import)"), "python"),
    (
        re.compile(r"((function\s+\w+\(|const|let|var)\s*.*?=>|\b(document|console|window)\.)"),
        "javascript",
    ),
    (re.compile(r"<(\w+)(.*?)>.*?</\1>", re.S), "html"),
    (re.compile(r"<\?php"), "php"),
    (re.compile(r"(public|private|protected)\s+(static\s+)?(void|int|String)"), "java"),
    (re.compile(r"#include\s+<"), "c/c++"),
    (re.compile(r"using\s+System;"), "c#"),
    (re.compile(r"(SELECT|INSERT|UPDATE|DELETE)\s+.*FROM", re.I), "sql"),
    (re.compile(r":\s+\w+\s*\{"), "css"),
    (re.compile(r"^[^:]+:\s* # YAML key-value", re.M | re.X), "yaml"),
    (re.compile(r"\$\w+"), "shell/bash"),
]

# Markdown processing regex
_BULLET_RX = re.compile(r"^[•‣▪◦‧﹒∙·] ?", re.MULTILINE)

# Lazy Loading State
_tiktoken_enc_instance: Union["Encoding", bool, None] = None

# OCR Caching (Simple in-memory - can be extended)
_OCR_CACHE: Dict[str, Any] = {}

# Domain Rules and Compiled Regex (Loaded Lazily)
_DOMAIN_RULES_CACHE: Optional[Dict] = None
_ACTIVE_DOMAIN: Optional[str] = None
_BOUND_RX: Optional[re.Pattern] = None
_CUSTOM_SECT_RX: Optional[List[Tuple[re.Pattern, str]]] = None
_METRIC_RX: Optional[List[Tuple[str, re.Pattern]]] = None
_RISK_RX: Optional[Dict[str, re.Pattern]] = None
_DOC_LABELS: Optional[List[str]] = None
_CLASS_PROMPT_PREFIX: Optional[str] = None

###############################################################################
# Utility & Helper Functions (Private Module Level)                           #
###############################################################################


def _log_dependency_warnings():
    """Logs warnings for missing optional dependencies on first use."""
    if not _DOCLING_AVAILABLE:
        logger.warning(
            "Docling library not available. Advanced PDF/Office conversion features disabled."
        )
    if not _PYPDF2_AVAILABLE:
        logger.warning("PyPDF2 not available. Basic PDF fallback conversion disabled.")
    if not _DOCX_AVAILABLE:
        logger.warning("python-docx not available. Basic DOCX fallback conversion disabled.")
    if not _PANDAS_AVAILABLE:
        logger.warning("Pandas not available. Pandas output format for tables disabled.")
    if not _TIKTOKEN_AVAILABLE:
        logger.warning(
            "Tiktoken not available. Token-based chunking will fallback to character chunking."
        )
    ocr_deps = {
        "Pillow": _PIL_AVAILABLE,
        "numpy": _NUMPY_AVAILABLE,
        "opencv-python": _CV2_AVAILABLE,
        "pytesseract": _PYTESSERACT_AVAILABLE,
        "pdf2image": _PDF2IMAGE_AVAILABLE,
    }
    missing_ocr = [name for name, avail in ocr_deps.items() if not avail]
    if missing_ocr:
        logger.warning(
            f"Missing OCR dependencies: {', '.join(missing_ocr)}. OCR functionality limited/disabled."
        )
    if not _PDFPLUMBER_AVAILABLE and not _PYMUPDF_AVAILABLE:
        logger.warning(
            "Missing direct PDF text extraction libraries (pdfplumber/pymupdf). Direct text extraction disabled."
        )
    elif not _PDFPLUMBER_AVAILABLE:
        logger.warning(
            "pdfplumber not available. Will rely solely on PyMuPDF for direct text extraction."
        )
    elif not _PYMUPDF_AVAILABLE:
        logger.warning(
            "PyMuPDF not available. Will rely solely on pdfplumber for direct text extraction."
        )
    if not _TRAFILATURA_AVAILABLE:
        logger.warning("Trafilatura not installed. Trafilatura HTML extraction disabled.")
    if not _READABILITY_AVAILABLE:
        logger.warning("Readability-lxml not installed. Readability HTML extraction disabled.")
    if not _markdownify_fallback:
        logger.warning("Markdownify not installed. HTML to Markdown fallback disabled.")


# Call once on import to log status
_log_dependency_warnings()


def _load_and_compile_domain_rules():
    """Loads domain rules from config and compiles regex patterns."""
    global _DOMAIN_RULES_CACHE, _ACTIVE_DOMAIN, _BOUND_RX, _CUSTOM_SECT_RX
    global _METRIC_RX, _RISK_RX, _DOC_LABELS, _CLASS_PROMPT_PREFIX

    if _DOMAIN_RULES_CACHE is not None:  # Already loaded
        return

    logger.debug("Lazily loading and compiling domain rules...")
    default_rules = {
        "generic": {
            "classification": {
                "labels": ["Report", "Contract", "Presentation", "Memo", "Email", "Manual"],
                "prompt_prefix": "Classify the document into exactly one of: ",
            },
            "sections": {
                "boundary_regex": r"^\s*(chapter\s+\d+|section\s+\d+|[A-Z][A-Za-z\s]{3,80})$",
                "custom": [],
            },
            "metrics": {
                "metric_1": {"aliases": ["metric one", "m1"]},
                "metric_2": {"aliases": ["metric two", "m2"]},
            },
            "risks": {"Risk_A": r"risk a", "Risk_B": r"risk b"},
        },
        "finance": {
            "classification": {
                "labels": [
                    "10-K",
                    "Credit Agreement",
                    "Investor Deck",
                    "Press Release",
                    "Board Minutes",
                    "NDA",
                    "LPA",
                    "CIM",
                ],
                "prompt_prefix": "Identify the document type (finance domain): ",
            },
            "sections": {
                "boundary_regex": r"^\s*(item\s+\d+[a-z]?\.|[A-Z][A-Za-z\s]{3,80})$",
                "custom": [
                    {"regex": r"item\s+1a?\.? .*business", "label": "Business"},
                    {"regex": r"item\s+1a\.? .*risk factors", "label": "Risk Factors"},
                    {"regex": r"item\s+7\.? .*management'?s discussion", "label": "MD&A"},
                    {"regex": r"covena[nv]ts", "label": "Covenants"},
                ],
            },
            "metrics": {
                "revenue": {
                    "aliases": ["revenue", "net sales", "total sales", "sales revenue", "turnover"]
                },
                "ebitda": {
                    "aliases": ["ebitda", "adj. ebitda", "operating profit", "operating income"]
                },
                "gross_profit": {"aliases": ["gross profit"]},
                "net_income": {"aliases": ["net income", "net profit", "earnings"]},
                "capex": {"aliases": ["capital expenditures", "capex"]},
                "debt": {"aliases": ["total debt", "net debt", "long-term debt"]},
            },
            "risks": {
                "Change_of_Control": r"change\s+of\s+control",
                "ESG_Risk": r"(child\s+labor|environmental\s+violation|scope\s+3)",
                "PII": r"(\bSSN\b|social security number|passport no)",
            },
        },
        "legal": {
            "classification": {
                "labels": ["Contract", "NDA", "Lease", "Policy", "License", "Settlement"],
                "prompt_prefix": "Classify the legal document into exactly one of: ",
            },
            "sections": {
                "boundary_regex": r"^\s*(article\s+\d+|section\s+\d+|[A-Z][A-Za-z\s]{3,80})$",
                "custom": [
                    {"regex": r"definitions", "label": "Definitions"},
                    {"regex": r"termination", "label": "Termination"},
                    {"regex": r"confidentiality", "label": "Confidentiality"},
                ],
            },
            "metrics": {},
            "risks": {
                "Indemnity": r"indemnif(y|ication)",
                "Liquidated_Damages": r"liquidated damages",
                "Governing_Law_NY": r"governing law.*new york",
                "Governing_Law_DE": r"governing law.*delaware",
            },
        },
        "medical": {
            "classification": {
                "labels": [
                    "Clinical Study",
                    "Patient Report",
                    "Lab Results",
                    "Prescription",
                    "Care Plan",
                ],
                "prompt_prefix": "Classify the medical document: ",
            },
            "sections": {
                "boundary_regex": r"^\s*(section\s+\d+|[A-Z][A-Za-z\s]{3,80})$",
                "custom": [
                    {"regex": r"diagnosis", "label": "Diagnosis"},
                    {"regex": r"treatment", "label": "Treatment"},
                    {"regex": r"medications", "label": "Medications"},
                    {"regex": r"allergies", "label": "Allergies"},
                ],
            },
            "metrics": {
                "blood_pressure": {"aliases": ["blood pressure", "bp"]},
                "heart_rate": {"aliases": ["heart rate", "hr"]},
                "temperature": {"aliases": ["temperature", "temp"]},
                "bmi": {"aliases": ["bmi", "body mass index"]},
            },
            "risks": {
                "Allergy": r"allergic reaction",
                "Contraindication": r"contraindicat(ed|ion)",
                "Adverse_Event": r"adverse event",
            },
        },
    }
    _DOMAIN_RULES_CACHE = default_rules

    _ACTIVE_DOMAIN = os.getenv("DOC_DOMAIN", "generic")
    # Config loading if needed:
    # from ultimate_mcp_server.config import get_config
    # try:
    #     cfg = get_config()
    #     _ACTIVE_DOMAIN = cfg.document_processing.domain if cfg and hasattr(cfg, 'document_processing') else "generic"
    # except Exception as e:
    #     logger.warning(f"Failed to load document processing domain from config: {e}. Defaulting to 'generic'.")
    #     _ACTIVE_DOMAIN = "generic"

    if _ACTIVE_DOMAIN not in _DOMAIN_RULES_CACHE:
        logger.warning(f"Unknown DOC_DOMAIN '{_ACTIVE_DOMAIN}', defaulting to 'generic'.")
        _ACTIVE_DOMAIN = "generic"

    instruction_json = _DOMAIN_RULES_CACHE[_ACTIVE_DOMAIN]

    try:
        _BOUND_RX = re.compile(instruction_json["sections"].get("boundary_regex", r"$^"), re.M)
    except re.error as e:
        logger.error(f"Invalid boundary regex for domain {_ACTIVE_DOMAIN}: {e}")
        _BOUND_RX = re.compile(r"$^")

    _CUSTOM_SECT_RX = []
    for d in instruction_json["sections"].get("custom", []):
        try:
            _CUSTOM_SECT_RX.append((re.compile(d["regex"], re.I), d["label"]))
        except re.error as e:
            logger.error(
                f"Invalid custom section regex '{d['regex']}' for domain {_ACTIVE_DOMAIN}: {e}"
            )

    _METRIC_RX = []
    for key, cfg in instruction_json.get("metrics", {}).items():
        aliases = cfg.get("aliases", [])
        if aliases:
            try:
                sorted_aliases = sorted(aliases, key=len, reverse=True)
                joined = "|".join(re.escape(a) for a in sorted_aliases)
                if joined:
                    pattern = re.compile(
                        rf"""(?i)\b({joined})\b[\s:–-]*([$€£]?\s?-?\d[\d,.]*)""",
                        re.VERBOSE | re.MULTILINE,
                    )
                    _METRIC_RX.append((key, pattern))
            except re.error as e:
                logger.error(
                    f"Invalid metric regex for alias group '{key}' in domain {_ACTIVE_DOMAIN}: {e}"
                )

    _RISK_RX = {}
    for t, pat_str in instruction_json.get("risks", {}).items():
        try:
            _RISK_RX[t] = re.compile(pat_str, re.I)
        except re.error as e:
            logger.error(
                f"Invalid risk regex for '{t}' in domain {_ACTIVE_DOMAIN}: '{pat_str}'. Error: {e}"
            )

    _DOC_LABELS = instruction_json["classification"].get("labels", [])
    _CLASS_PROMPT_PREFIX = instruction_json["classification"].get("prompt_prefix", "")
    logger.info(f"Domain rules loaded and compiled for domain: '{_ACTIVE_DOMAIN}'")


def _get_active_domain_rules():
    """Ensures domain rules are loaded and returns them."""
    if _DOMAIN_RULES_CACHE is None:
        _load_and_compile_domain_rules()
    return {
        "active_domain": _ACTIVE_DOMAIN,
        "bound_rx": _BOUND_RX,
        "custom_sect_rx": _CUSTOM_SECT_RX,
        "metric_rx": _METRIC_RX,
        "risk_rx": _RISK_RX,
        "doc_labels": _DOC_LABELS,
        "class_prompt_prefix": _CLASS_PROMPT_PREFIX,
    }


def _get_tiktoken_encoder() -> Optional["Encoding"]:
    """Lazy load and return the tiktoken encoder instance."""
    global _tiktoken_enc_instance
    if _tiktoken_enc_instance is not None:
        return (
            _tiktoken_enc_instance
            if isinstance(_tiktoken_enc_instance, tiktoken.Encoding)
            else None
        )
    if not _TIKTOKEN_AVAILABLE:
        _tiktoken_enc_instance = False
        return None
    try:
        encoding_name = os.getenv("TIKTOKEN_ENCODING", "cl100k_base")
        logger.info(f"Lazy-loading tiktoken encoding: {encoding_name}")
        _tiktoken_enc_instance = tiktoken.get_encoding(encoding_name)  # type: ignore
        logger.info("Successfully lazy-loaded tiktoken encoder.")
        return _tiktoken_enc_instance  # type: ignore
    except Exception as e:
        logger.error(f"Failed to lazy-load tiktoken: {e}", exc_info=True)
        _tiktoken_enc_instance = False
        return None


async def _standalone_llm_call(
    *,
    prompt: str,
    provider: str = Provider.OPENAI.value,
    model: str | None = None,
    temperature: float = 0.3,
    max_tokens: int | None = None,
    extra: Dict[str, Any] | None = None,
) -> str:
    """Standalone wrapper to make LLM calls using the completion tool."""
    if not callable(generate_completion):
        logger.error("LLM generation function 'generate_completion' is not available.")
        raise ToolError("LLM_UNAVAILABLE", details={"reason": "generate_completion not available"})

    chosen_provider = provider
    try:
        additional_params = extra or {}
        response_dict = await generate_completion(
            prompt=prompt,
            provider=chosen_provider,
            model=model,
            temperature=temperature,
            max_tokens=max_tokens,
            additional_params=additional_params,
        )
        if isinstance(response_dict, dict):
            if response_dict.get("isError", False) or not response_dict.get("success", True):
                err_detail = response_dict.get("error", {})
                err_msg = err_detail.get("message", "Unknown LLM Error")
                err_code = err_detail.get("type", "LLM_CALL_FAILED")
                logger.error(
                    f"LLM call failed [{err_code}]: {err_msg}. Raw Response: {response_dict}"
                )
                raise ToolError(
                    err_code,
                    details={
                        "provider": chosen_provider,
                        "error": err_msg,
                        "raw_response": str(response_dict),
                    },
                )
            llm_content = response_dict.get("text") or response_dict.get("content")
            if llm_content is None:
                logger.error(f"LLM response missing 'text'/'content': {response_dict}")
                raise ToolError(
                    "LLM_INVALID_RESPONSE",
                    details={"reason": "Missing content", "response_received": str(response_dict)},
                )
            if isinstance(llm_content, str):
                return llm_content.strip()
            else:
                logger.warning(f"LLM content not string: {type(llm_content)}. Converting.")
                return str(llm_content).strip()
        else:
            logger.error(f"LLM response unexpected format: {response_dict}")
            raise ToolError(
                "LLM_INVALID_RESPONSE", details={"response_received": str(response_dict)}
            )
    except ProviderError as pe:
        logger.error(f"LLM provider error ({chosen_provider}): {pe}", exc_info=True)
        raise ToolError(
            "LLM_PROVIDER_ERROR",
            details={"provider": chosen_provider, "error_code": pe.error_code, "error": str(pe)},
        ) from pe
    except ToolError as te:
        raise te
    except Exception as e:
        logger.error(f"LLM call failed ({chosen_provider}): {e}", exc_info=True)
        raise ToolError(
            "LLM_CALL_FAILED", details={"provider": chosen_provider, "error": str(e)}
        ) from e


@contextmanager
def _span(label: str):
    """Context manager for timing operations (module level)."""
    st = time.perf_counter()
    logger.debug(f"Starting span: {label}")
    try:
        yield
    finally:
        elapsed = time.perf_counter() - st
        logger.debug(f"Finished span: {label} ({elapsed:.3f}s)")


def _get_docling_converter(device, threads: int):
    """Create a Docling DocumentConverter."""
    if not _DOCLING_AVAILABLE:
        raise ToolError("DEPENDENCY_MISSING", details={"dependency": "docling"})
    if (
        not PdfPipelineOptions
        or not AcceleratorOptions
        or not InputFormat
        or not PdfFormatOption
        or not DocumentConverter
    ):
        raise ToolError(
            "INTERNAL_ERROR", details={"reason": "Docling partially imported but types missing"}
        )
    opts = PdfPipelineOptions()
    opts.do_ocr = False
    opts.generate_page_images = False
    opts.do_table_extraction = True # Explicitly enable table extraction in the pipeline options
    opts.accelerator_options = AcceleratorOptions(num_threads=threads, device=device)
    try:
        converter_options = {InputFormat.PDF: PdfFormatOption(pipeline_options=opts)}
        return DocumentConverter(format_options=converter_options)
    except Exception as e:
        logger.error(f"Failed to initialize Docling DocumentConverter: {e}", exc_info=True)
        raise ToolError(
            "INITIALIZATION_FAILED", details={"component": "DocumentConverter", "error": str(e)}
        ) from e


def _get_input_path_or_temp(
    document_path: Optional[str], document_data: Optional[bytes]
) -> Tuple[Path, bool]:
    """Gets a valid Path object for input. Saves data to temp file if needed."""
    is_temp = False
    if document_path:
        path = Path(document_path)
        if not path.is_file():
            raise ToolInputError(
                f"Input file not found: {document_path}", param_name="document_path"
            )
        return path, is_temp
    elif document_data:
        try:
            suffix = ".bin"
            if document_data.startswith(b"%PDF"):
                suffix = ".pdf"
            elif len(document_data) > 10 and document_data[6:10] in (b"JFIF", b"Exif"):
                suffix = ".jpg"
            elif document_data.startswith(b"\x89PNG\r\n\x1a\n"):
                suffix = ".png"
            elif document_data.startswith((b"II*\x00", b"MM\x00*")):
                suffix = ".tiff"
            elif document_data.startswith(b"PK\x03\x04"):
                suffix = ".zip"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(document_data)
                path = Path(tmp_file.name)
            is_temp = True
            logger.debug(f"Saved input data to temporary file: {path}")
            return path, is_temp
        except Exception as e:
            raise ToolError(
                "TEMP_FILE_ERROR",
                details={"error": f"Failed to save input data to temporary file: {e}"},
            ) from e
    else:
        raise ToolInputError("Either 'document_path' or 'document_data' must be provided.")


@contextmanager
def _handle_temp_file(path: Path, is_temp: bool):
    """Context manager to clean up temporary file."""
    try:
        yield path
    finally:
        if is_temp and path.exists():
            try:
                path.unlink()
                logger.debug(f"Cleaned up temporary file: {path}")
            except OSError as e:
                logger.warning(f"Failed to delete temporary file {path}: {e}")


def _tmp_path(src: str, fmt: str) -> Path:
    """Generate a temporary file path for output."""
    src_path = Path(src.split("?")[0])
    stem = src_path.stem or "document"
    ext = "md" if fmt == "markdown" else fmt
    timestamp = int(time.time() * 1000)
    temp_dir = Path(tempfile.gettempdir())
    temp_dir.mkdir(parents=True, exist_ok=True)
    return temp_dir / f"{stem}_{timestamp}.{ext}"


def _get_docling_metadata(doc: Any) -> dict[str, Any]:
    """Extract metadata from a Docling document."""
    if not _DOCLING_AVAILABLE or not doc:
        return {"error": "Docling not available or document object missing"}
    num_pages = 0
    try:
        num_pages = doc.num_pages() if callable(getattr(doc, "num_pages", None)) else 0
        has_tables = False
        has_figures = False
        has_sections = False
        if hasattr(doc, "pages") and isinstance(doc.pages, list):
            for page in doc.pages:
                if hasattr(page, "content") and page.content:
                    if (
                        callable(getattr(page.content, "has_tables", None))
                        and page.content.has_tables()
                    ):
                        has_tables = True
                    if (
                        callable(getattr(page.content, "has_figures", None))
                        and page.content.has_figures()
                    ):
                        has_figures = True
                if has_tables and has_figures:
                    break
        if hasattr(doc, "texts") and isinstance(doc.texts, list):
            for item in doc.texts:
                if hasattr(item, "__class__") and item.__class__.__name__ == "SectionHeaderItem":
                    has_sections = True
                    break
                elif hasattr(item, "label") and getattr(item, "label", None) == "section_header":
                    has_sections = True
                    break
        return {
            "num_pages": num_pages,
            "has_tables": has_tables,
            "has_figures": has_figures,
            "has_sections": has_sections,
        }
    except Exception as e:
        logger.warning(f"Docling metadata collection failed: {e}", exc_info=True)
        return {
            "num_pages": num_pages,
            "has_tables": False,
            "has_figures": False,
            "has_sections": False,
            "metadata_error": str(e),
        }


def _get_basic_metadata(text_content: str, num_pages: int = 0) -> dict[str, Any]:
    """Generate basic metadata for non-Docling content."""
    has_tables = "| --- |" in text_content or "\t" in text_content
    has_figures = "![" in text_content
    has_sections = bool(re.search(r"^#{1,6}\s+", text_content, re.M))
    return {
        "num_pages": num_pages,
        "has_tables": has_tables,
        "has_figures": has_figures,
        "has_sections": has_sections,
    }


def _json(obj: Any) -> str:
    """Utility to serialize objects to JSON."""
    return json.dumps(obj, ensure_ascii=False, separators=(",", ":"))


def _hash(txt: str) -> str:
    """Generate SHA-1 hash of text."""
    return hashlib.sha1(txt.encode("utf-8", "ignore")).hexdigest()


# --- HTML Helpers ---
def _is_html_fragment(text: str) -> bool:
    """Check if text contains likely HTML markup using precompiled patterns."""
    check_len = min(len(text), 5000)
    sample = text[:check_len]
    return any(p.search(sample) for p in _HTML_PATTERNS)


def _best_soup(html_txt: str) -> Tuple[Any, str]:
    """Try progressively more forgiving parsers; fall back to empty soup."""
    BeautifulSoup, Tag = _get_bs4()
    
    parsers_to_try = ["html.parser", "lxml", "html5lib"]
    for p_name in parsers_to_try:
        try:
            return BeautifulSoup(html_txt, p_name), p_name
        except Exception as e_parse:
            logger.debug(f"Parser {p_name} failed on full HTML: {e_parse}")
            continue

    try:
        # Last-ditch effort: wrap in minimal html/body structure and try basic parser
        wrapped_html = f"<html><body>{html_txt}</body></html>"
        return BeautifulSoup(wrapped_html, "html.parser"), "html.parser-fragment"
    except Exception as e_frag:
        logger.warning(
            f"Fragment parsing also failed: {e_frag}. Returning empty soup.", exc_info=True
        )
        return BeautifulSoup("", "html.parser"), "failed"


def _clean_html(html_txt: str) -> Tuple[str, str]:
    """Cleans potentially unsafe HTML elements/attributes; returns (cleaned_html, parser_used)."""
    BeautifulSoup, Tag = _get_bs4()
    
    try:
        soup, parser_used = _best_soup(html_txt)
        
        # Remove potentially problematic tags
        tags_to_remove = [
            "script",
            "style", 
            "noscript",
            "iframe",
            "object",
            "embed",
            "applet",
            "form",
            "input",
            "button",
            "textarea",
            "select",
            "option",
            "meta",
            "link",
        ]
        for el in soup(tags_to_remove):
            el.decompose()

        # Clean dangerous attributes
        dangerous_attrs = ["onclick", "onload", "onerror", "href", "src", "action"]
        for tag in soup.find_all(True):
            current_attrs = list(tag.attrs.keys())
            for attr in current_attrs:
                attr_val_str = str(tag.get(attr, "")).lower()
                is_unsafe = (
                    attr.lower() in dangerous_attrs
                    or "javascript:" in attr_val_str
                    or "data:" in attr_val_str
                    or "vbscript:" in attr_val_str
                )
                if is_unsafe and attr in tag.attrs:
                    del tag[attr]

        text = str(soup)
        logger.debug(f"HTML cleaned with parser: {parser_used}")
        return str(soup), parser_used
    except Exception as e:
        try:
            return str(soup), parser_used
        except Exception as stringify_error:
            logger.error(f"Could not stringify soup after error: {stringify_error}")
            return html_txt, "fallback"  # Return original if all else fails


# --- Markdown Helpers ---
def _sanitize(md: str) -> str:
    """Basic Markdown sanitization."""
    if not md:
        return ""
    md = md.replace("\u00a0", " ")
    md = _BULLET_RX.sub("- ", md)
    md = re.sub(r"\n{3,}", "\n\n", md)
    md = re.sub(r"[ \t]+$", "", md, flags=re.MULTILINE)
    md = re.sub(r"^[ \t]+", "", md, flags=re.MULTILINE)
    md = re.sub(r"(^|\n)(#{1,6})([^#\s])", r"\1\2 \3", md)
    md = re.sub(r"```\s*\n", "```\n", md)
    md = re.sub(r"\n\s*```", "\n```", md)
    md = re.sub(r"^[*+]\s", "- ", md, flags=re.MULTILINE)
    md = re.sub(r"^\d+\.\s", lambda m: f"{m.group(0).strip()} ", md, flags=re.MULTILINE)
    return md.strip()


def _improve(md: str) -> str:
    """Apply structural improvements to Markdown text."""
    if not md:
        return ""
    # Ensure blank lines around major block elements
    md = re.sub(r"(?<=\S)\n(#{1,6}\s)", r"\n\n\1", md)
    md = re.sub(r"(^#{1,6}\s.*\S)\n(?!\n|#|```|>|\s*[-*+]|\s*\d+\.)", r"\1\n\n", md, flags=re.M)
    md = re.sub(r"(?<=\S)\n(```)", r"\n\n\1", md)
    md = re.sub(r"(```)\n(?!\n)", r"\1\n\n", md)
    md = re.sub(r"(?<=\S)\n(> )", r"\n\n\1", md)
    md = re.sub(r"(\n> .*\S)\n(?!\n|>\s)", r"\1\n\n", md, flags=re.M)
    md = re.sub(r"(?<=\S)\n(\s*([-*+]|\d+\.)\s)", r"\n\n\1", md)
    md = re.sub(
        r"(\n(\s*[-*+]\s+|\s*\d+\.\s+).*\S)\n(?!\n|\s*([-*+]|\d+\.)\s)", r"\1\n\n", md, flags=re.M
    )
    md = re.sub(r"(?<=\S)\n(-{3,}|\*{3,}|_{3,})$", r"\n\n\1", md, flags=re.M)
    md = re.sub(r"(^-{3,}|\*{3,}|_{3,})\n(?!\n)", r"\1\n\n", md, flags=re.M)
    md = re.sub(r"\n{3,}", "\n\n", md)
    return md.strip()


def _convert_html_table_to_markdown(table_tag: Any) -> str:
    """Converts a single BeautifulSoup table Tag to a Markdown string."""
    md_rows = []

    # Try to find header from thead first
    header_row_tag = table_tag.find("thead")
    header_cells_tags = []
    if header_row_tag:
        header_cells_tags = header_row_tag.find_all(["th", "td"], recursive=False)
        if not header_cells_tags:
            header_row_tr = header_row_tag.find("tr")
            if header_row_tr:
                header_cells_tags = header_row_tr.find_all(["th", "td"])
    if not header_cells_tags:
        first_row = table_tag.find("tr")
        if first_row:
            temp_cells = first_row.find_all(["th", "td"])
            is_header = any(c.name == "th" for c in temp_cells) or (
                len(temp_cells) > 0
                and not any(re.match(r"^\s*[\d.,-]+\s*$", c.get_text()) for c in temp_cells)
            )
            if is_header:
                header_cells_tags = temp_cells

    if header_cells_tags:
        num_cols = len(header_cells_tags)
        hdr = [
            " ".join(c.get_text(" ", strip=True).replace("|", "\\|").split())
            for c in header_cells_tags
        ]
        md_rows.append("| " + " | ".join(hdr) + " |")
        md_rows.append("| " + " | ".join(["---"] * num_cols) + " |")
    else:
        body_rows_tags = (
            table_tag.find("tbody").find_all("tr")
            if table_tag.find("tbody")
            else table_tag.find_all("tr")
        )
        if not body_rows_tags:
            return ""
        for r in body_rows_tags:
            num_cols = max(num_cols, len(r.find_all(["th", "td"])))
        if num_cols == 0:
            return ""
        logger.debug(f"Table has no clear header, assuming {num_cols} columns.")
        md_rows.append("| " + " | ".join([f"Col {i + 1}" for i in range(num_cols)]) + " |")
        md_rows.append("| " + " | ".join(["---"] * num_cols) + " |")

    body_rows_tags = []
    tbody = table_tag.find("tbody")
    if tbody:
        body_rows_tags = tbody.find_all("tr")
    else:
        all_trs = table_tag.find_all("tr")
        start_index = (
            1
            if header_cells_tags
            and all_trs
            and header_cells_tags[0].find_parent("tr") == all_trs[0]
            else 0
        )
        body_rows_tags = all_trs[start_index:]

    for r in body_rows_tags:
        cells = r.find_all(["td", "th"])
        cell_texts = [
            " ".join(c.get_text(" ", strip=True).replace("|", "\\|").split()) for c in cells
        ]
        cell_texts.extend([""] * (num_cols - len(cells)))
        cell_texts = cell_texts[:num_cols]
        md_rows.append("| " + " | ".join(cell_texts) + " |")

    return "\n".join(md_rows)


def _convert_html_tables_to_markdown(html_txt: str) -> str:
    """Convert HTML tables to Markdown format."""
    BeautifulSoup, Tag = _get_bs4()
    
    try:
        soup, parser_used = _best_soup(html_txt)

        # Find all tables and convert them
        tables = soup.find_all("table")
        if not tables:
            return str(soup)

        for table_tag in tables:
            try:
                md_table_str = _convert_html_table_to_markdown(table_tag)
                # Replace the table with the markdown version
                placeholder = soup.new_string(f"\n\n{md_table_str}\n\n")
                table_tag.replace_with(placeholder)
                # Clean up the old table
                table_tag.decompose()
            except Exception as e:
                logger.warning(f"Failed to convert table to markdown: {e}")

        return str(soup)
    except Exception as e:
        logger.error(f"Error converting HTML tables: {e}")
        return html_txt


def _html_to_md_core(html_txt: str, links: bool, imgs: bool, tbls: bool, width: int) -> str:
    """Convert HTML to Markdown using primary and fallback libraries."""
    try:
        html2text_module = _get_html2text()
        if html2text_module is None:
            raise ImportError("html2text module not available")
        
        h = html2text_module.HTML2Text()
        h.ignore_links = not links
        h.ignore_images = not imgs
        processed_html = html_txt
        if tbls:
            processed_html = _convert_html_tables_to_markdown(html_txt)
            h.ignore_tables = True
        else:
            h.ignore_tables = True

        h.body_width = width if width > 0 else 0
        h.unicode_snob = True
        h.escape_snob = True
        h.skip_internal_links = True
        h.single_line_break = True

        md_text = h.handle(processed_html)
        logger.debug("html2text conversion successful.")
        return md_text.strip()
    except Exception as e_html2text:
        logger.warning(f"html2text failed ({e_html2text}); attempting fallback with markdownify")
        # Markdownify fallback is not available, use simple text extraction
        try:
            BeautifulSoup, Tag = _get_bs4()
            soup, _ = _best_soup(processed_html)
            # Simple text extraction as fallback
            text = soup.get_text(separator=' ', strip=True)
            return text
        except Exception as e_fallback:
            logger.error(f"All HTML to markdown conversion methods failed: {e_fallback}", exc_info=True)
            raise ToolError(
                "MARKDOWN_CONVERSION_FAILED",
                details={
                    "reason": "All conversion methods failed", 
                    "html2text_error": str(e_html2text),
                    "fallback_error": str(e_fallback),
                },
            ) from e_html2text


###############################################################################
# Core OCR & PDF Helper Functions (Standalone)                                #
###############################################################################


def _ocr_check_dep(dep_name: str, is_available: bool, feature: str):
    """Checks if a required dependency is available, raising ToolError if not."""
    if not is_available:
        logger.error(f"Missing required dependency '{dep_name}' for feature '{feature}'.")
        raise ToolError("DEPENDENCY_MISSING", details={"dependency": dep_name, "feature": feature})


def _ocr_extract_text_from_pdf_direct(
    file_path: Path, start_page: int = 0, max_pages: int = 0
) -> Tuple[List[str], bool]:
    """
    Extracts text directly from PDF using PyMuPDF or PDFPlumber (sync function).

    Args:
        file_path: Path to the PDF file.
        start_page: 0-based starting page index.
        max_pages: Maximum number of pages to extract (0 for all from start_page).

    Returns:
        Tuple containing:
        - List of strings, one per extracted page (or error marker).
        - Boolean indicating if meaningful text was found on at least one page.
    """
    texts: List[str] = []
    has_text = False
    min_chars = 50  # Threshold for considering a page to have meaningful text

    # --- Try PyMuPDF First ---
    if _PYMUPDF_AVAILABLE and pymupdf:
        logger.debug(f"Attempting direct text extraction with PyMuPDF for {file_path}")
        try:
            with pymupdf.open(file_path) as doc:  # type: ignore
                total_pages = len(doc)
                # Calculate 0-based end page index (exclusive)
                end_page = (
                    total_pages if max_pages <= 0 else min(start_page + max_pages, total_pages)
                )
                # Ensure start_page is valid
                start_page = min(start_page, total_pages)
                end_page = max(start_page, end_page)  # Ensure end is not before start

                for i in range(start_page, end_page):
                    try:
                        page = doc.load_page(i)  # Use load_page for clarity
                        page_text = page.get_text("text") or ""  # Specify text format
                        texts.append(page_text)
                        if len(page_text.strip()) >= min_chars:
                            has_text = True
                    except Exception as e_page:
                        logger.warning(
                            f"PyMuPDF: Error extracting text from page {i + 1}: {e_page}"
                        )
                        texts.append(f"[Page {i + 1} Extraction Error: PyMuPDF]")
                logger.debug(
                    f"PyMuPDF extracted {len(texts)} pages. Found meaningful text: {has_text}"
                )
                return texts, has_text
        except Exception as e_pymupdf:
            logger.warning(
                f"PyMuPDF direct text extraction failed: {e_pymupdf}. Trying PDFPlumber..."
            )
            # Fall through to PDFPlumber if PyMuPDF failed

    # --- Try PDFPlumber as Fallback ---
    if _PDFPLUMBER_AVAILABLE and pdfplumber:
        logger.debug(f"Attempting direct text extraction with PDFPlumber for {file_path}")
        try:
            # pdfplumber might require explicit closing
            pdf = pdfplumber.open(file_path)  # type: ignore
            try:
                total_pages = len(pdf.pages)
                end_page = (
                    total_pages if max_pages <= 0 else min(start_page + max_pages, total_pages)
                )
                start_page = min(start_page, total_pages)
                end_page = max(start_page, end_page)

                for i in range(start_page, end_page):
                    try:
                        page = pdf.pages[i]
                        # Use slightly more tolerant settings
                        page_text = (
                            page.extract_text(x_tolerance=2, y_tolerance=2, keep_blank_chars=True)
                            or ""
                        )
                        texts.append(page_text)
                        if len(page_text.strip()) >= min_chars:
                            has_text = True
                    except Exception as e_page:
                        logger.warning(
                            f"PDFPlumber: Error extracting text from page {i + 1}: {e_page}"
                        )
                        texts.append(f"[Page {i + 1} Extraction Error: PDFPlumber]")
                logger.debug(
                    f"PDFPlumber extracted {len(texts)} pages. Found meaningful text: {has_text}."
                )
                return texts, has_text
            finally:
                pdf.close()  # Ensure file handle is closed
        except Exception as e_plumber:
            logger.error(f"PDFPlumber direct text extraction failed: {e_plumber}", exc_info=True)
            # If PyMuPDF also failed (or wasn't available), raise the final error
            if (
                not _PYMUPDF_AVAILABLE
            ):  # Only raise if it was the only option tried or PyMuPDF failed before
                raise ToolError(
                    "DIRECT_EXTRACTION_FAILED",
                    details={"reason": "PDFPlumber failed", "error": str(e_plumber)},
                ) from e_plumber
            else:  # PyMuPDF failed first, now PDFPlumber failed
                raise ToolError(
                    "DIRECT_EXTRACTION_FAILED",
                    details={
                        "reason": "Both PyMuPDF and PDFPlumber failed",
                        "error": str(e_plumber),
                    },
                ) from e_plumber

    # --- If neither library worked ---
    logger.error(
        "No functional direct PDF text extraction library (PyMuPDF or PDFPlumber) available or both failed."
    )
    raise ToolError("DIRECT_EXTRACTION_FAILED", details={"reason": "No available/working library"})


def _ocr_convert_pdf_to_images(
    file_path: Path, start_page: int = 0, max_pages: int = 0, dpi: int = 300
) -> List["PILImage.Image"]:
    """Converts PDF path pages to PIL Images using 0-based indexing internally (sync function)."""
    _ocr_check_dep("pdf2image", _PDF2IMAGE_AVAILABLE, "PDF->Image Conversion")
    _ocr_check_dep("Pillow", _PIL_AVAILABLE, "PDF->Image Conversion")
    if convert_from_path is None:
        raise ToolError("INTERNAL_ERROR", details={"reason": "pdf2image.convert_from_path is None"})

    try:
        # pdf2image uses 1-based indexing for first_page/last_page args
        first_page_1based = start_page + 1
        last_page_1based = None if max_pages <= 0 else first_page_1based + max_pages - 1
        logger.debug(
            f"Converting PDF {file_path} (pages {first_page_1based}-{last_page_1based or 'end'}, dpi={dpi})"
        )

        with _span(f"pdf2image_path_p{first_page_1based}-{last_page_1based or 'end'}"):
            # pdf2image handles its own temporary files internally if output_folder=None
            # Using a TemporaryDirectory might be slightly less efficient but ensures cleanup
            with tempfile.TemporaryDirectory() as temp_dir:
                images = convert_from_path(  # type: ignore
                    file_path,
                    dpi=dpi,
                    first_page=first_page_1based,
                    last_page=last_page_1based,
                    output_folder=temp_dir,  # Recommended for multi-threading stability
                    fmt="png",
                    thread_count=max(1, os.cpu_count() // 2 if os.cpu_count() else 1),
                    use_pdftocairo=True,  # Often more reliable than pdftoppm
                )
        logger.info(f"Converted {len(images)} pages from PDF path.")
        return images  # type: ignore
    except Exception as e:  # Catch specific pdf2image errors if library defines them
        logger.error(f"PDF path to image conversion failed: {e}", exc_info=True)
        raise ToolError(
            "PDF_CONVERSION_FAILED", details={"reason": "pdf2image path failed", "error": str(e)}
        ) from e


def _ocr_preprocess_image(
    image: "PILImage.Image", preprocessing_options: Optional[Dict[str, Any]] = None
) -> "PILImage.Image":
    """Preprocesses an image for better OCR results (sync function)."""
    if not _PIL_AVAILABLE:
        logger.warning("Pillow (PIL) not available. Skipping preprocessing.")
        return image
    if not ImageEnhance or not ImageFilter:  # Check specifically for submodules
        logger.warning("PIL ImageEnhance or ImageFilter missing. Some enhancements skipped.")

    can_use_cv2 = _CV2_AVAILABLE and _NUMPY_AVAILABLE and cv2 is not None and np is not None
    if (
        not can_use_cv2
        and preprocessing_options
        and any(k in preprocessing_options for k in ["denoise", "threshold", "deskew"])
    ):
        logger.warning("OpenCV/NumPy missing. Advanced preprocessing disabled.")

    prep_opts = {
        "denoise": True,
        "threshold": "otsu",
        "deskew": True,
        "enhance_contrast": True,
        "enhance_brightness": False,
        "enhance_sharpness": False,
        "apply_filters": [],
        "resize_factor": 1.0,
        **(preprocessing_options or {}),
    }
    logger.debug(f"Applying preprocessing with options: {prep_opts}")

    img_pil = image.copy()
    # Apply PIL enhancements first
    if ImageEnhance:
        if prep_opts.get("enhance_brightness"):
            img_pil = ImageEnhance.Brightness(img_pil).enhance(1.3)
        if prep_opts.get("enhance_contrast") and not can_use_cv2:
            img_pil = ImageEnhance.Contrast(img_pil).enhance(1.4)
        if prep_opts.get("enhance_sharpness"):
            img_pil = ImageEnhance.Sharpness(img_pil).enhance(1.5)
    if ImageFilter:
        filters = prep_opts.get("apply_filters", [])
        for filter_name in filters:
            try:
                if filter_name == "unsharp_mask":
                    img_pil = img_pil.filter(ImageFilter.UnsharpMask(radius=2, percent=150))
                elif filter_name == "detail":
                    img_pil = img_pil.filter(ImageFilter.DETAIL)
                elif filter_name == "edge_enhance":
                    img_pil = img_pil.filter(ImageFilter.EDGE_ENHANCE)
                elif filter_name == "smooth":
                    img_pil = img_pil.filter(ImageFilter.SMOOTH)
                else:
                    logger.warning(f"Unknown PIL filter: {filter_name}")
            except Exception as e:
                logger.warning(f"PIL filter '{filter_name}' failed: {e}")

    if not can_use_cv2:
        return img_pil  # Return PIL-enhanced if CV2 unavailable

    # OpenCV Processing
    try:
        img_cv = np.array(img_pil)
        if len(img_cv.shape) == 3 and img_cv.shape[2] == 3:
            gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)
        elif len(img_cv.shape) == 3 and img_cv.shape[2] == 4:
            gray = cv2.cvtColor(img_cv, cv2.COLOR_RGBA2GRAY)
        else:
            gray = img_cv

        original_height, original_width = gray.shape[:2]
        deskewed_gray = gray.copy()  # Operate on this copy

        # Deskewing (best on grayscale before thresholding might change shapes)
        if prep_opts.get("deskew", True):
            try:
                # Use inverted image for finding text blocks if background is light
                mean_intensity = np.mean(gray)
                invert_for_deskew = mean_intensity > 128
                deskew_input = cv2.bitwise_not(gray) if invert_for_deskew else gray

                # Use a less aggressive threshold for finding angle
                _, angle_thresh = cv2.threshold(
                    deskew_input, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
                )
                coords = cv2.findNonZero(angle_thresh)
                if coords is not None and len(coords) > 10:
                    angle = cv2.minAreaRect(coords)[-1]
                    if angle < -45:
                        angle = -(90 + angle)
                    else:
                        angle = -angle

                    if abs(angle) > 0.1:
                        (h, w) = gray.shape[:2]
                        center = (w // 2, h // 2)
                        M = cv2.getRotationMatrix2D(center, angle, 1.0)
                        # Rotate original grayscale image
                        deskewed_gray = cv2.warpAffine(
                            gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE
                        )
                        logger.debug(f"Deskewed image by {angle:.2f} degrees.")
            except Exception as e_deskew:
                logger.warning(f"Deskewing failed: {e_deskew}. Using original orientation.")
                deskewed_gray = gray  # Reset to original gray if deskew fails

        processed_img = deskewed_gray  # Start processing from (potentially) deskewed gray

        # Adaptive scaling calculation (applied later)
        resize_factor = prep_opts.get("resize_factor", 1.0)
        if resize_factor == 1.0:
            longest_edge = max(original_width, original_height)
            target_low, target_high = 1500, 3500
            if 0 < longest_edge < target_low:
                resize_factor = math.ceil(target_low / longest_edge * 10) / 10
            elif longest_edge > target_high:
                resize_factor = math.floor(target_high / longest_edge * 10) / 10
            resize_factor = max(0.5, min(3.0, resize_factor))

        # Contrast enhancement on grayscale
        if prep_opts.get("enhance_contrast", True):
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            processed_img = clahe.apply(processed_img)

        # Denoising grayscale (before thresholding)
        if prep_opts.get("denoise", True):
            # Adjust h based on image size? Might be overkill.
            # h_param = math.ceil(10 * math.log10(max(10, min(original_width, original_height))))
            processed_img = cv2.fastNlMeansDenoising(processed_img, None, 10, 7, 21)

        # Thresholding
        threshold_method = prep_opts.get("threshold", "otsu")
        if threshold_method == "otsu":
            # No need for blur if denoised already
            _, processed_img = cv2.threshold(
                processed_img, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
            )
        elif threshold_method == "adaptive":
            block_size = max(11, math.floor(min(processed_img.shape[:2]) / 20) * 2 + 1)
            processed_img = cv2.adaptiveThreshold(
                processed_img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, block_size, 5
            )
        # If no threshold, check background and invert if needed for Tesseract
        elif np.mean(processed_img) < 128:
            processed_img = cv2.bitwise_not(processed_img)

        # Resizing (applied last)
        if resize_factor != 1.0:
            current_h, current_w = processed_img.shape[:2]
            new_w = math.ceil(current_w * resize_factor)
            new_h = math.ceil(current_h * resize_factor)
            processed_img = cv2.resize(processed_img, (new_w, new_h), interpolation=cv2.INTER_CUBIC)
            logger.debug(f"Resized image by factor {resize_factor:.2f} to {new_w}x{new_h}")

        final_pil_image = Image.fromarray(processed_img)
        return final_pil_image
    except Exception as e_cv:
        logger.error(f"OpenCV preprocessing failed: {e_cv}", exc_info=True)
        return img_pil  # Fallback to PIL-processed image


def _ocr_run_tesseract(
    image: "PILImage.Image", ocr_language: str = "eng", ocr_config: str = ""
) -> str:
    """Extracts text from an image using Tesseract OCR (sync function)."""
    _ocr_check_dep("pytesseract", _PYTESSERACT_AVAILABLE, "OCR Text Extraction")
    _ocr_check_dep("Pillow", _PIL_AVAILABLE, "OCR Text Extraction")
    if pytesseract is None:
        raise ToolError("INTERNAL_ERROR", details={"reason": "pytesseract is None"})
    try:
        # Combine language and custom config flags
        custom_config = f"-l {ocr_language} {ocr_config}".strip()
        logger.debug(f"Running Tesseract with config: '{custom_config}'")
        with _span(f"pytesseract_ocr_{ocr_language}"):
            # Use a timeout? Tesseract can sometimes hang. Requires subprocess handling.
            # For simplicity, no timeout implemented here.
            text = pytesseract.image_to_string(
                image, config=custom_config, timeout=60
            )  # Add 60s timeout
        logger.debug(f"Tesseract extracted {len(text)} characters.")
        return text or ""  # Ensure string return
    except pytesseract.TesseractNotFoundError as e:
        logger.error("Tesseract executable not found or not in PATH.")
        raise ToolError("DEPENDENCY_MISSING", details={"dependency": "Tesseract OCR Engine"}) from e
    except RuntimeError as e_runtime:  # Catch Tesseract runtime errors (like timeout)
        logger.error(f"Tesseract runtime error: {e_runtime}", exc_info=True)
        raise ToolError(
            "OCR_FAILED", details={"engine": "Tesseract", "error": f"Runtime error: {e_runtime}"}
        ) from e_runtime
    except Exception as e:
        logger.error(f"Tesseract OCR extraction failed: {e}", exc_info=True)
        raise ToolError("OCR_FAILED", details={"engine": "Tesseract", "error": str(e)}) from e


def _ocr_is_text_mostly_noise(text: str, noise_threshold: float = 0.4) -> bool:
    """
    Determine if extracted text is mostly noise based on character distribution.
    Considers alphanumeric, whitespace, and common punctuation as 'valid'.

    Args:
        text: The text string to analyze.
        noise_threshold: The ratio (0.0 to 1.0) of non-valid characters above which
                         the text is considered noisy. Default is 0.4 (40%).

    Returns:
        True if the text is considered mostly noise, False otherwise.
    """
    if not text or not isinstance(text, str):
        return False  # Empty or invalid input is not noise

    text_length = len(text)
    if text_length < 20:  # Don't evaluate very short strings
        return False

    # Define a set of characters generally expected in non-noisy text
    # (alphanumeric, whitespace, common punctuation/symbols)
    # Adding more symbols that might appear legitimately in documents
    valid_char_pattern = re.compile(r"[a-zA-Z0-9\s.,;:!?\"'()\[\]{}%/$£€¥₽₹#@&*+=<>~|_^-]")

    valid_chars_count = len(valid_char_pattern.findall(text))

    # Calculate the ratio of characters *not* matching the valid pattern
    noise_ratio = 1.0 - (valid_chars_count / text_length)

    is_noise = noise_ratio > noise_threshold
    if is_noise:
        # Log only a snippet to avoid flooding logs with potentially large noisy text
        snippet = text.replace("\n", " ")[:100]  # Replace newlines for cleaner log output
        logger.debug(
            f"Text flagged as noisy (Ratio: {noise_ratio:.2f} > {noise_threshold}): '{snippet}...'"
        )

    return is_noise


def _ocr_is_likely_header_or_footer(text: str, line_length_threshold: int = 80) -> bool:
    """
    Determine if a single line of text is likely a header or footer based on common patterns.

    Args:
        text: The line of text to evaluate.
        line_length_threshold: Lines longer than this are less likely to be headers/footers. Default 80.

    Returns:
        True if the line matches common header/footer patterns, False otherwise.
    """
    text = text.strip()
    if not text or len(text) > line_length_threshold:
        return False

    # --- Pattern Checks ---
    # 1. Page Number patterns (robust check)
    #    - "Page X", "P. X", "X / Y", "- X -", etc.
    #    - Allows for variations in spacing and separators
    if re.search(r"(?i)\b(page|p[ag]{1,2}\.?|seite|s\.?)\s*\d+", text):
        return True
    if re.match(r"^\s*[-–—]?\s*\d+\s*[/of\s]+\s*\d+\s*[-–—]?\s*$", text):
        return True  # e.g., "1 / 10", "1 of 10"
    if re.match(r"^\s*[-–—]?\s*\d+\s*[-–—]?\s*$", text):
        return True  # Just a number, possibly bracketed

    # 2. Date patterns
    #    - "Month Day, Year", "DD/MM/YYYY", "YYYY-MM-DD", etc.
    if re.search(
        r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}", text, re.I
    ):
        return True
    if re.search(r"\b\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\b", text):
        return True
    if re.search(r"\b\d{4}[./-]\d{1,2}[./-]\d{1,2}\b", text):
        return True  # ISO-like

    # 3. Common keywords (case-insensitive start of line)
    if re.match(
        r"^(confidential|internal use only|draft|proprietary|for discussion purposes)", text, re.I
    ):
        return True
    if re.match(r"^(copyright|\(c\)|©)\s*\d*", text, re.I):
        return True

    # 4. Repeated characters (often used as separators)
    #    - Check if the line consists mostly of one or two non-alphanumeric characters
    non_alnum_chars = re.sub(r"[a-zA-Z0-9\s]", "", text)
    if len(non_alnum_chars) > 5 and len(set(non_alnum_chars)) <= 2:
        return True

    # 5. Company Names / Document Titles (Heuristic - might be too broad)
    #    - Check if it's short, title-cased, and doesn't end in punctuation?
    # if len(text.split()) < 7 and text == text.title() and not text.endswith(('.', '?', '!')):
    #     # Further check: Is this text repeated elsewhere? (Needs broader context)
    #     pass # This heuristic is often unreliable without more context.

    # 6. All Caps Short Lines (Potential titles/headers)
    if text.isupper() and len(text.split()) < 7 and len(text) > 3:
        return True

    return False  # Default: Not a header/footer


def _ocr_remove_headers_and_footers(text: str, max_lines_check: int = 5) -> str:
    """
    Removes likely headers and footers from the top/bottom of the text block.

    Args:
        text: The block of text (potentially multiple pages concatenated).
        max_lines_check: How many lines from the top and bottom to examine. Default 5.

    Returns:
        The text with potential header/footer lines removed.
    """
    if not text or not isinstance(text, str):
        return ""

    lines = text.splitlines()
    num_lines = len(lines)

    # Don't process if text is too short to reliably identify headers/footers
    if num_lines < max_lines_check * 2:
        return text

    lines_to_remove_indices: Set[int] = set()

    # Check top lines
    for i in range(max_lines_check):
        if i < num_lines:  # Ensure index is valid
            line_text = lines[i]
            # Also check if the line is very short (e.g., just whitespace remnants)
            if _ocr_is_likely_header_or_footer(line_text) or len(line_text.strip()) <= 2:
                lines_to_remove_indices.add(i)
            # Stop checking top lines if a probable content line is found early
            elif len(line_text) > 80 and i < max_lines_check // 2:  # Heuristic for content line
                break
        else:  # Should not happen given initial num_lines check, but safety
            break

    # Check bottom lines
    for i in range(max_lines_check):
        idx = num_lines - 1 - i
        # Ensure index is valid and not already marked for removal from top scan
        if idx >= 0 and idx not in lines_to_remove_indices:
            line_text = lines[idx]
            if _ocr_is_likely_header_or_footer(line_text) or len(line_text.strip()) <= 2:
                lines_to_remove_indices.add(idx)
            # Stop checking bottom lines if a probable content line is found early
            elif len(line_text) > 80 and i < max_lines_check // 2:
                break
        elif idx < 0:  # Reached top of file during bottom check
            break

    if not lines_to_remove_indices:
        return text  # No lines identified for removal

    logger.debug(f"Removing {len(lines_to_remove_indices)} potential header/footer lines.")

    # Build the result, skipping removed lines
    result_lines = [line for i, line in enumerate(lines) if i not in lines_to_remove_indices]

    # Remove leading/trailing blank lines potentially left after removal
    # This needs care: find first/last non-blank line indices
    first_content_line = -1
    last_content_line = -1
    for i, line in enumerate(result_lines):
        if line.strip():
            if first_content_line == -1:
                first_content_line = i
            last_content_line = i

    if first_content_line == -1:  # All lines were removed or blank
        return ""
    else:
        # Join only the content lines, preserving internal blank lines
        cleaned_text = "\n".join(result_lines[first_content_line : last_content_line + 1])
        return cleaned_text


async def _ocr_enhance_text_chunk(
    chunk: str, output_format: str = "markdown", remove_headers: bool = False
) -> str:
    """Enhances OCR text chunk using LLM (standalone internal helper)."""
    # --- Apply Basic Rule-based Cleaning First ---
    cleaned_text = chunk.strip()  # Work on a copy

    # Join words incorrectly split across lines (common OCR artifact)
    cleaned_text = re.sub(r"(\w+)-\s*\n\s*(\w+)", r"\1\2", cleaned_text)

    # Normalize multiple whitespace characters (including newlines within paragraphs)
    # This is aggressive and might merge intended line breaks within code/poetry
    # Consider a less aggressive approach if preserving specific line breaks is crucial.
    # cleaned_text = re.sub(r"\s+", " ", cleaned_text) # Too aggressive

    # Normalize space/tab characters to single space
    cleaned_text = re.sub(r"[ \t]+", " ", cleaned_text)
    # Collapse multiple blank lines (2+ newlines) into exactly two newlines
    cleaned_text = re.sub(r"\n\s*\n", "\n\n", cleaned_text)

    # Optional header/footer removal using rules *before* LLM
    if remove_headers:
        original_len = len(cleaned_text)
        cleaned_text = _ocr_remove_headers_and_footers(cleaned_text)
        if len(cleaned_text) < original_len:
            logger.debug("Applied rule-based header/footer removal pre-LLM.")

    # Check for noise after initial cleaning
    if _ocr_is_text_mostly_noise(cleaned_text):
        logger.warning(
            "Text chunk noisy after basic cleaning, LLM enhancement might be less effective."
        )
        # Decide whether to proceed or return early based on noise level?
        # For now, proceed with LLM enhancement.

    # --- LLM Prompt Generation ---
    format_instruction = ""
    if output_format == "markdown":
        format_instruction = """
2. Format as clean, readable markdown:
   - Use appropriate heading levels (#, ##, etc.). Infer structure where possible.
   - Format lists correctly (bulleted or numbered).
   - Apply emphasis (*italic*) and strong (**bold**) sparingly where appropriate.
   - Represent tabular data using markdown table syntax IF table structure is clearly identifiable.
   - Use code blocks (```) for code snippets or equations if detected."""
    else:  # output_format == "text"
        format_instruction = """
2. Format as clean, readable plain text:
   - Ensure clear paragraph separation (double newline).
   - Maintain list structures with standard markers (e.g., -, 1.).
   - Avoid markdown syntax like #, *, _, ```, etc."""

    header_footer_instruction = (
        "Remove any remaining headers, footers, and page numbers."
        if remove_headers
        else "Preserve all content including potential headers/footers."
    )
    prompt = f"""You are an expert text processor specialized in correcting OCR errors from scanned documents. Please process the following text according to these instructions:

1. Fix OCR-induced errors:
   - Correct character recognition errors (e.g., 'rn' vs 'm', 'O' vs '0', 'l' vs '1', 'S' vs '5').
   - Join words incorrectly split across lines (e.g., "hyphen-\nation").
   - Merge paragraphs that were artificially split by page breaks or scanning artifacts.
   - Split run-on paragraphs where a clear topic shift or structural break (like a list starting) occurs.
   - Use context to resolve ambiguities and reconstruct the original meaning accurately.
{format_instruction}
3. Clean up formatting:
   - Remove redundant spaces within lines.
   - Ensure consistent paragraph spacing (double newline between paragraphs).
   - {header_footer_instruction}

4. IMPORTANT: Preserve all meaningful content and the original structure as much as possible. Do not add information or summaries. Do not change the substance of the text. Focus solely on fixing OCR errors and applying the requested formatting based *only* on the input text provided.

Input Text:
```text
{cleaned_text}
```

Corrected Output ({output_format}):"""

    try:
        logger.debug(
            f"Sending chunk (len={len(cleaned_text)}) to LLM for enhancement (format={output_format}, rm_hdrs={remove_headers})."
        )
        # Use a capable model (adjust model name as needed)
        provider = Provider.OPENAI.value
        model = "gpt-4o-mini"

        # Estimate max tokens needed
        estimated_input_tokens = len(cleaned_text) // 3
        buffer_factor = 1.4 if output_format == "markdown" else 1.2  # Slightly more buffer
        llm_max_tokens = int(estimated_input_tokens * buffer_factor) + 500
        # Cap based on typical context window limits (e.g., ~16k tokens for GPT-4 Turbo input, allow ample output)
        llm_max_tokens = max(1000, min(llm_max_tokens, 8000))

        # Assume _standalone_llm_call is defined elsewhere
        enhanced_text = await _standalone_llm_call(
            prompt=prompt,
            provider=provider,
            model=model,
            temperature=0.1,  # Very low temperature for factual correction
            max_tokens=llm_max_tokens,
        )

        # --- Post-processing LLM Output ---
        # Remove potential preamble/apologies
        enhanced_text = re.sub(
            r"^(Okay, |Here is |Sure, |Here['']s |Certainly, |Based on the text provided.*?\n)[:\n]?\s*",
            "",
            enhanced_text,
            flags=re.IGNORECASE | re.DOTALL,
        )
        # Remove potential markdown fences around the whole output
        enhanced_text = re.sub(
            r"^\s*```(?:\w+\n)?([\s\S]*?)\n?```\s*$", r"\1", enhanced_text
        ).strip()

        logger.debug(f"LLM enhancement returned text (len={len(enhanced_text)}).")
        return enhanced_text

    except ToolError as e:
        # Log the specific ToolError and fallback
        logger.error(
            f"LLM text enhancement failed with ToolError: {e.error_code} - {str(e)}. Returning pre-LLM cleaned text."
        )
        return cleaned_text
    except Exception as e:
        # Log unexpected errors and fallback
        logger.error(f"Unexpected error during LLM text enhancement: {e}", exc_info=True)
        return cleaned_text


def _ocr_validate_file_path(file_path: str, expected_extension: Optional[str] = None) -> Path:
    """Validates a file path exists and optionally has the expected extension."""
    if not file_path or not isinstance(file_path, str):
        raise ToolInputError("File path cannot be empty or non-string", param_name="file_path")

    try:
        # Expand user directory and normalize path separators
        path = Path(os.path.expanduser(os.path.normpath(file_path)))
    except Exception as e:
        raise ToolInputError(
            f"Invalid file path format: {file_path}. Error: {e}", param_name="file_path"
        ) from e

    if not path.exists():
        raise ToolInputError(f"File not found at path: {path}", param_name="file_path")
    if not path.is_file():
        raise ToolInputError(f"Path exists but is not a file: {path}", param_name="file_path")
    # Check extension case-insensitively
    if expected_extension and not path.suffix.lower() == expected_extension.lower():
        raise ToolInputError(
            f"File does not have the expected extension ({expected_extension}): {path}",
            param_name="file_path",
        )
    # Optional: Check read permissions?
    # if not os.access(path, os.R_OK):
    #     raise ToolInputError(f"Cannot read file (permission denied): {path}", param_name="file_path")

    return path


def _ocr_detect_tables(image: "PILImage.Image") -> List[Tuple[int, int, int, int]]:
    """Detects potential tables in an image using OpenCV (sync function)."""
    # Check dependencies first
    if not _CV2_AVAILABLE or not _NUMPY_AVAILABLE or not _PIL_AVAILABLE:
        logger.warning("Cannot detect tables: OpenCV, NumPy, or Pillow not available.")
        return []
    # Ensure library objects are valid
    if cv2 is None or np is None:
        logger.warning("Cannot detect tables: OpenCV or NumPy object is None.")
        return []

    try:
        img = np.array(image)
        # Convert to grayscale if necessary
        if len(img.shape) == 3:
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        elif len(img.shape) == 2:
            gray = img
        else:
            logger.warning(f"Unexpected image shape for table detection: {img.shape}")
            return []

        # --- Table Detection Logic (Example using line detection) ---
        # 1. Thresholding (Adaptive often works well for lines)
        thresh_inv = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 15, 5
        )

        # 2. Detect Horizontal Lines
        horizontal_kernel = cv2.getStructuringElement(
            cv2.MORPH_RECT, (min(40, gray.shape[1] // 10), 1)
        )  # Kernel size relative to width
        detected_horizontal = cv2.morphologyEx(
            thresh_inv, cv2.MORPH_OPEN, horizontal_kernel, iterations=2
        )
        cnts_h, _ = cv2.findContours(
            detected_horizontal, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE
        )

        # 3. Detect Vertical Lines
        vertical_kernel = cv2.getStructuringElement(
            cv2.MORPH_RECT, (1, min(40, gray.shape[0] // 10))
        )  # Kernel size relative to height
        detected_vertical = cv2.morphologyEx(
            thresh_inv, cv2.MORPH_OPEN, vertical_kernel, iterations=2
        )
        cnts_v, _ = cv2.findContours(detected_vertical, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        # 4. Combine contours or find bounding boxes of large contours containing lines
        # Strategy: Find large contours in the original inverted threshold image,
        # then check if those contours contain significant horiz/vert lines.
        contours, _ = cv2.findContours(thresh_inv, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        table_regions = []
        img_area = img.shape[0] * img.shape[1]
        min_table_area = img_area * 0.01  # Lower threshold slightly (1%)
        min_dimension = 50  # Min width/height for a contour to be considered

        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            area = w * h
            aspect_ratio = w / max(1, h)

            # Basic filtering based on size and aspect ratio
            if (
                area > min_table_area
                and w > min_dimension
                and h > min_dimension
                and 0.1 < aspect_ratio < 10.0
            ):
                # Check for significant presence of detected lines within this bounding box
                roi_h = detected_horizontal[y : y + h, x : x + w]
                roi_v = detected_vertical[y : y + h, x : x + w]
                # Heuristic: Check if non-zero pixels (lines) exceed a small fraction of the ROI area or length
                min_line_pixels_h = w * 0.3  # Require horizontal lines covering ~30% width
                min_line_pixels_v = h * 0.3  # Require vertical lines covering ~30% height
                if (
                    cv2.countNonZero(roi_h) > min_line_pixels_h
                    and cv2.countNonZero(roi_v) > min_line_pixels_v
                ):
                    table_regions.append((x, y, w, h))
                # else:
                #    logger.debug(f"Contour rejected: area={area}, w={w}, h={h}, h_px={cv2.countNonZero(roi_h)}, v_px={cv2.countNonZero(roi_v)}")

        # Optional: Merge overlapping bounding boxes (omitted for simplicity)
        # merged_regions = merge_overlapping_boxes(table_regions) # Needs implementation

        logger.debug(f"Detected {len(table_regions)} potential table regions.")
        return table_regions

    except Exception as e:
        logger.error(f"OpenCV Table detection failed: {e}", exc_info=True)
        return []


def _ocr_process_toc(toc: List) -> List[Dict[str, Any]]:
    """Processes a PDF table of contents (from PyMuPDF) into a nested structure."""
    if not toc:
        return []
    result: List[Dict[str, Any]] = []
    # Stack stores tuples: (level, parent_list_to_append_to)
    stack: List[Tuple[int, List]] = [(-1, result)]
    for item in toc:
        # PyMuPDF TOC item format: [level, title, page, ?dest_dict]
        if not isinstance(item, (list, tuple)) or len(item) < 3:
            logger.warning(f"Skipping malformed TOC item: {item}")
            continue
        try:
            level = int(item[0])
            title = str(item[1])
            page = int(item[2])
        except (ValueError, TypeError, IndexError) as e:
            logger.warning(f"Error parsing TOC item '{item}': {e}")
            continue

        # Pop stack until parent level is found
        while stack[-1][0] >= level:
            stack.pop()
            if not stack:  # Should not happen with initial (-1, result)
                logger.error("TOC stack became empty unexpectedly.")
                return result  # Return what we have so far

        # Create new entry and add to parent's children list
        entry: Dict[str, Any] = {"title": title, "page": page, "children": []}
        stack[-1][1].append(entry)
        # Push current entry onto stack for potential children
        stack.append((level, entry["children"]))
    return result


def _ocr_split_text_into_chunks(
    text: str, max_chunk_size: int = 8000, overlap: int = 200
) -> List[str]:
    """Splits text into chunks, trying to respect paragraphs and sentences (sync function)."""
    if not text or not isinstance(text, str):
        return []

    max_chunk_size = max(1000, min(max_chunk_size, 15000))  # Sensible limits
    overlap = max(50, min(overlap, max_chunk_size // 4))
    # Ensure min_chunk_size is reasonable, at least larger than overlap
    min_chunk_size = max(overlap * 2, 100)

    chunks = []
    start_index = 0
    text_len = len(text)

    while start_index < text_len:
        end_index = min(start_index + max_chunk_size, text_len)

        # Handle the last chunk directly
        if end_index == text_len:
            chunk = text[start_index:end_index]
            # Only add if it has meaningful content (more than just whitespace)
            if chunk.strip():
                chunks.append(chunk)
            break  # End of text reached

        best_split_index = -1
        # Prefer double newline (paragraph break)
        split_point_para = text.rfind("\n\n", max(start_index, end_index - overlap * 2), end_index)
        if split_point_para != -1 and split_point_para > start_index:  # Ensure split is after start
            # Check if this split results in a reasonably sized chunk
            if (split_point_para + 2 - start_index) >= min_chunk_size:
                best_split_index = split_point_para + 2

        # If no good paragraph break, try sentence breaks
        if best_split_index == -1:
            sentence_break_pattern = r"[.?!]['\"]?(\s|\n|$)"  # Include end of string
            # Search within a reasonable lookback window
            search_region_start = max(start_index, end_index - overlap)
            search_region = text[search_region_start:end_index]
            matches = list(re.finditer(sentence_break_pattern, search_region))
            if matches:
                # Find the offset of the last match within the search region
                last_match_end_offset = matches[-1].end()
                # Calculate the split point relative to the original string
                split_point_sentence = search_region_start + last_match_end_offset
                # Check if this split is valid and creates a reasonably sized chunk
                if (
                    split_point_sentence > start_index
                    and (split_point_sentence - start_index) >= min_chunk_size
                ):
                    best_split_index = split_point_sentence

        # Fallback to single newline or space if still no good break
        if best_split_index == -1:
            split_point_newline = text.rfind("\n", max(start_index, end_index - overlap), end_index)
            split_point_space = text.rfind(" ", max(start_index, end_index - overlap), end_index)
            # Choose the latest valid break (newline or space)
            split_point_fallback = max(split_point_newline, split_point_space)
            if (
                split_point_fallback > start_index
                and (split_point_fallback + 1 - start_index) >= min_chunk_size
            ):
                best_split_index = split_point_fallback + 1

        # Force split at max_chunk_size boundary if no suitable break found,
        # or if the best found break is too early (making the chunk too small)
        if (
            best_split_index <= start_index
            or (best_split_index - start_index) < min_chunk_size // 2
        ):
            # Check if simply taking end_index results in a valid chunk start for next iteration
            potential_next_start = max(start_index + 1, end_index - overlap)
            if potential_next_start < text_len:  # Avoid forcing if it's the last chunk anyway
                best_split_index = end_index
            else:  # If forcing split here would make the loop end, try a slightly earlier hard split?
                # For simplicity, let's stick to end_index, the loop termination handles the last part.
                best_split_index = end_index

        # Extract the chunk
        chunk = text[start_index:best_split_index]
        if chunk.strip():  # Only add non-empty chunks
            chunks.append(chunk)

        # Calculate the start index for the next chunk
        next_start = max(start_index + 1, best_split_index - overlap)

        # Ensure substantial forward progress to avoid infinite loops on edge cases
        # Use max_chunk_size here instead of the undefined 'size'
        min_progress = min(max_chunk_size // 10, 50)  # Ensure we advance by at least a small amount
        next_start = max(next_start, start_index + min_progress)

        # Safety check: don't let next_start go beyond the text length
        start_index = min(next_start, text_len)

    # Filter out any potential empty strings added during edge cases
    final_chunks = [c for c in chunks if c]

    logger.debug(f"Split text ({text_len} chars) into {len(final_chunks)} chunks")
    return final_chunks


async def _ocr_assess_text_quality(original_text: str, enhanced_text: str) -> Dict[str, Any]:
    """Assesses the quality of OCR enhancement using LLM (Standalone)."""
    if not original_text and not enhanced_text:
        return {"score": 0, "explanation": "No text provided for assessment.", "examples": []}
    if not original_text:
        return {
            "score": 100,
            "explanation": "Original text was empty, enhanced text provided.",
            "examples": [],
        }
    if not enhanced_text:
        return {
            "score": 0,
            "explanation": "Enhanced text is empty, original text was not.",
            "examples": [],
        }

    max_sample = 4000
    original_sample = original_text[:max_sample] + (
        "..." if len(original_text) > max_sample else ""
    )
    enhanced_sample = enhanced_text[:max_sample] + (
        "..." if len(enhanced_text) > max_sample else ""
    )

    prompt = f"""Please assess the quality improvement from the 'Original OCR Text' to the 'Enhanced Text'. Focus on:
1. Correction of OCR errors (typos, spacing, broken words).
2. Improvement in formatting and readability (paragraphs, lists, structure).
3. Accuracy in preserving the original meaning and content.
4. Effectiveness of removing noise (like headers/footers if applicable).

Original OCR Text:
```
{original_sample}
```

Enhanced Text:
```
{enhanced_sample}
```

Provide your assessment ONLY in the following JSON format:
{{
  "score": <integer score 0-100, where 100 is perfect enhancement>,
  "explanation": "<brief explanation of the score, highlighting key improvements or remaining issues>",
  "examples": [
    "<example 1 of a specific correction or improvement>",
    "<example 2>",
    "<example 3 (optional)>"
  ]
}}
Do not add any text before or after the JSON object.
"""

    try:
        logger.debug("Requesting LLM quality assessment.")
        assessment_json_str = await _standalone_llm_call(
            prompt=prompt, max_tokens=500, temperature=0.2
        )
        try:
            json_match = re.search(r"```(?:json)?\s*([\s\S]+?)\s*```", assessment_json_str)
            json_str = json_match.group(1).strip() if json_match else assessment_json_str.strip()
            start_brace = json_str.find("{")
            end_brace = json_str.rfind("}")
            if start_brace != -1 and end_brace != -1 and start_brace < end_brace:
                json_str = json_str[start_brace : end_brace + 1]
            elif not json_str.startswith("{"):
                raise ValueError("Could not find JSON object boundaries.")

            assessment_data = json.loads(json_str)
            if (
                not isinstance(assessment_data, dict)
                or "score" not in assessment_data
                or "explanation" not in assessment_data
                or "examples" not in assessment_data
                or not isinstance(assessment_data["examples"], list)
            ):
                raise ValueError("Parsed JSON has incorrect structure.")
            try:
                assessment_data["score"] = (
                    int(assessment_data["score"]) if assessment_data["score"] is not None else None
                )
            except (ValueError, TypeError):
                assessment_data["score"] = None
            assessment_data["explanation"] = str(assessment_data["explanation"])
            assessment_data["examples"] = [str(ex) for ex in assessment_data["examples"]]
            logger.debug(
                f"Quality assessment received: Score {assessment_data.get('score', 'N/A')}"
            )
            return assessment_data
        except (json.JSONDecodeError, ValueError, TypeError) as e:
            logger.error(
                f"Failed to parse quality assessment JSON: {e}. Raw:\n{assessment_json_str}"
            )
            return {
                "score": None,
                "explanation": f"Parse failed: {e}",
                "examples": [],
                "raw_response": assessment_json_str,
            }
    except Exception as e:
        logger.error(f"Error during LLM quality assessment call: {e}", exc_info=True)
        return {"score": None, "explanation": f"LLM call failed: {e}", "examples": []}


# --- Fallback Conversion Helpers (module level) ---
async def _fallback_convert_pdf(file_path: Path) -> Dict[str, Any]:
    """Basic PDF conversion using PyPDF2."""
    _ocr_check_dep("PyPDF2", _PYPDF2_AVAILABLE, "Basic PDF Fallback Conversion")
    try:
        logger.info(f"Using PyPDF2 fallback for PDF: {file_path}")
        content = ""
        metadata: Dict[str, Any] = {"is_fallback": True}
        num_pages = 0
        if PyPDF2 is None:
            raise ImportError("PyPDF2 object is None despite _PYPDF2_AVAILABLE=True")
        with open(file_path, "rb") as f:
            try:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                metadata["num_pages"] = num_pages
                pages = []
                for i in range(num_pages):
                    try:
                        page_text = reader.pages[i].extract_text() or ""
                        pages.append(page_text)
                    except Exception as page_err:
                        logger.warning(
                            f"PyPDF2 failed to extract text from page {i + 1}: {page_err}"
                        )
                        pages.append(f"[Page {i + 1} Extraction Error]")
                content = "\n\n".join(pages)
            except PyPDF2.errors.PdfReadError as pdf_err:
                logger.error(f"PyPDF2 could not read PDF {file_path}: {pdf_err}")
                raise ToolError(
                    "PDF_READ_ERROR", details={"library": "PyPDF2", "error": str(pdf_err)}
                ) from pdf_err
        metadata.update(_get_basic_metadata(content, num_pages))
        return {"content": content, "metadata": metadata}
    except Exception as e:
        logger.error(f"PyPDF2 fallback failed unexpectedly: {e}", exc_info=True)
        raise ToolError(
            "CONVERSION_FAILED",
            details={"file": str(file_path), "method": "PyPDF2 Fallback", "error": str(e)},
        ) from e


async def _fallback_convert_docx(file_path: Path) -> Dict[str, Any]:
    """Basic DOCX conversion using python-docx."""
    _ocr_check_dep("python-docx", _DOCX_AVAILABLE, "DOCX Fallback Conversion")
    try:
        logger.info(f"Using python-docx fallback for DOCX: {file_path}")
        if docx is None:
            raise ImportError("docx object is None despite _DOCX_AVAILABLE=True")
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text]
        content = "\n\n".join(paragraphs)
        metadata: Dict[str, Any] = {
            "num_pages": 0,
            "has_tables": len(doc.tables) > 0,
            "has_figures": len(doc.inline_shapes) > 0,
            "has_sections": len(doc.sections) > 0,
            "is_fallback": True,
        }
        metadata.update(_get_basic_metadata(content))
        return {"content": content, "metadata": metadata}
    except Exception as e:
        logger.error(f"python-docx fallback failed: {e}", exc_info=True)
        raise ToolError(
            "CONVERSION_FAILED",
            details={"file": str(file_path), "method": "python-docx Fallback", "error": str(e)},
        ) from e


async def _fallback_convert_text(file_path: Path) -> Dict[str, Any]:
    """Simple text file reading."""
    try:
        logger.info(f"Reading text file directly: {file_path}")
        content = file_path.read_text(encoding="utf-8", errors="replace")
        line_count = content.count("\n") + 1
        page_estimate = max(1, int(line_count / 50))
        metadata = {"num_pages": page_estimate, "is_fallback": True}
        metadata.update(_get_basic_metadata(content, page_estimate))
        return {"content": content, "metadata": metadata}
    except Exception as e:
        logger.error(f"Text file reading failed: {e}", exc_info=True)
        raise ToolError(
            "CONVERSION_FAILED",
            details={"file": str(file_path), "method": "Direct Text Read", "error": str(e)},
        ) from e


###############################################################################
# Standalone Tool Functions (Exportable)                                      #
###############################################################################


# ------------------------ Document Conversion -----------------------------
@with_tool_metrics
@with_error_handling
async def convert_document(
    document_path: Optional[str] = None,
    document_data: Optional[bytes] = None,
    output_format: str = "markdown",
    extraction_strategy: str = DEFAULT_EXTRACTION_STRATEGY,
    enhance_with_llm: bool = True,
    ocr_options: Optional[Dict] = None,
    output_path: Optional[str] = None,
    save_to_file: bool = False,
    page_range: Optional[str] = None,
    section_filter: Optional[str] = None,
    accelerator_device: str = "auto",
    num_threads: int = 4,
) -> Dict[str, Any]:
    """
    Convert documents (PDF, Office formats, Images) to various formats (Standalone Function).
    (Args/Returns docs same as original class method)
    """
    t0 = time.time()
    strategy = extraction_strategy.lower()
    output_format = output_format.lower()
    ocr_options = ocr_options or {}

    # --- Input Validation ---
    if not document_path and not document_data:
        raise ToolInputError("Either 'document_path' or 'document_data' must be provided.")
    if document_path and document_data:
        raise ToolInputError("Provide either 'document_path' or 'document_data', not both.")
    if strategy not in _VALID_EXTRACTION_STRATEGIES:
        raise ToolInputError(
            f"Invalid extraction_strategy. Choose from: {', '.join(_VALID_EXTRACTION_STRATEGIES)}",
            param_name="extraction_strategy",
            provided_value=strategy,
        )
    if output_format not in _VALID_FORMATS:
        raise ToolInputError(
            f"Invalid output_format. Choose from: {', '.join(_VALID_FORMATS)}",
            param_name="output_format",
            provided_value=output_format,
        )

    # --- Dependency Checks based on strategy ---
    if strategy == "docling":
        _ocr_check_dep("docling", _DOCLING_AVAILABLE, "Docling extraction strategy")
    if strategy in ["direct_text", "hybrid_direct_ocr"]:
        if not (_PYMUPDF_AVAILABLE or _PDFPLUMBER_AVAILABLE):
            raise ToolError(
                "DEPENDENCY_MISSING",
                details={
                    "dependency": "PyMuPDF or PDFPlumber",
                    "feature": "Direct Text strategy",
                },
            )
    if strategy in ["ocr", "hybrid_direct_ocr"]:
        _ocr_check_dep("pdf2image", _PDF2IMAGE_AVAILABLE, "OCR strategy")
        _ocr_check_dep("pytesseract", _PYTESSERACT_AVAILABLE, "OCR strategy")
        _ocr_check_dep("Pillow", _PIL_AVAILABLE, "OCR strategy")
        if ocr_options.get("preprocessing") and not (_CV2_AVAILABLE and _NUMPY_AVAILABLE):
            logger.warning(
                "Preprocessing options provided but OpenCV/NumPy missing. Preprocessing limited."
            )

    # Adjust output format compatibility
    effective_output_format = output_format
    if strategy != "docling" and output_format not in _OCR_COMPATIBLE_FORMATS:
        logger.warning(
            f"Output format '{output_format}' is not directly supported by strategy '{strategy}'. Defaulting to 'markdown'."
        )
        effective_output_format = "markdown"

    # --- Prepare Input ---
    input_path_obj: Optional[Path] = None
    is_temp_file = False
    input_name: str = "input_data"

    try:
        input_path_obj, is_temp_file = _get_input_path_or_temp(document_path, document_data)
        input_name = input_path_obj.name

        with _handle_temp_file(input_path_obj, is_temp_file) as current_input_path:
            input_suffix = current_input_path.suffix.lower()
            is_pdf = input_suffix == ".pdf"
            is_image = input_suffix in [".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"]
            is_office = input_suffix in [
                ".docx",
                ".pptx",
                ".xlsx",
                ".zip",
            ]  # Treat zip as potential office
            is_text = input_suffix in [".txt", ".md", ".html", ".xml", ".json"]

            # Validate strategy vs input type & adjust strategy if needed
            if not is_pdf and strategy in ["direct_text", "hybrid_direct_ocr"]:
                if is_image:
                    logger.warning(f"Strategy '{strategy}' needs PDF. Input is image. Using 'ocr'.")
                    strategy = "ocr"
                else:
                    raise ToolInputError(
                        f"Strategy '{strategy}' requires PDF input, got '{input_suffix}'."
                    )
            if not is_pdf and not is_image and strategy == "ocr":
                raise ToolInputError(
                    f"OCR strategy needs PDF/Image, got '{input_suffix}'. Use 'docling' or handle as text."
                )
            if is_office and strategy != "docling":
                if input_suffix == ".docx" and _DOCX_AVAILABLE:
                    logger.warning("Input is DOCX without 'docling'. Using fallback.")
                    strategy = "fallback_docx"
                # Add other office fallbacks here if needed
                else:
                    raise ToolInputError(
                        f"Office file ('{input_suffix}') requires 'docling' strategy or specific fallback library."
                    )
            if is_text and strategy != "docling":
                logger.info(f"Input is text ('{input_suffix}'). Using direct text handling.")
                strategy = "fallback_text"

            # --- Parse Page Range ---
            pages_to_process: Optional[List[int]] = None
            total_doc_pages = 0
            if page_range:
                try:
                    pages_set: Set[int] = set()
                    parts = page_range.split(",")
                    for part in parts:
                        part = part.strip()
                        if "-" in part:
                            start_str, end_str = part.split("-", 1)
                            start, end = int(start_str), int(end_str)
                            if start < 1 or end < start:
                                raise ValueError(f"Invalid range: {start}-{end}")
                            pages_set.update(range(start - 1, end))
                        else:
                            page_num = int(part)
                            if page_num < 1:
                                raise ValueError(f"Page number must be positive: {page_num}")
                            pages_set.add(page_num - 1)
                    if not pages_set:
                        raise ValueError("No valid pages selected.")
                    pages_to_process = sorted(list(pages_set))
                    logger.debug(
                        f"Parsed page range: {page_range} -> 0-based indices: {pages_to_process}"
                    )
                except ValueError as e:
                    raise ToolInputError(
                        f"Invalid page_range format: '{page_range}'. Error: {e}",
                        param_name="page_range",
                    ) from e

            # --- Result Structure Defaults ---
            result_content: Union[str, Dict] = ""
            doc_metadata: Dict[str, Any] = {}
            raw_text_pages: List[str] = []
            final_raw_text: Optional[str] = None
            quality_metrics: Optional[Dict] = None
            strategy_used = strategy

            # ======================== EXTRACTION STRATEGIES ========================

            if strategy == "docling":
                logger.info(f"Using 'docling' strategy for {input_name}")
                _ocr_check_dep("docling", _DOCLING_AVAILABLE, "Docling strategy")
                device_str = accelerator_device.lower()
                if device_str not in _ACCEL_MAP:
                    logger.warning(f"Invalid device '{device_str}', using 'auto'.")
                    device_str = "auto"
                device = _ACCEL_MAP[device_str]
                conv = _get_docling_converter(device, num_threads)
                loop = asyncio.get_running_loop()
                with _span("docling_conversion"):
                    docling_result = await loop.run_in_executor(
                        None, conv.convert, current_input_path
                    )
                if not docling_result or not docling_result.document:
                    raise ToolError("CONVERSION_FAILED", details={"reason": "Docling empty result"})
                doc_obj = docling_result.document
                doc_metadata = _get_docling_metadata(doc_obj)
                total_doc_pages = doc_metadata.get("num_pages", 0)

                if effective_output_format == "markdown":
                    result_content = doc_obj.export_to_markdown()
                elif effective_output_format == "text":
                    result_content = doc_obj.export_to_text()
                elif effective_output_format == "html":
                    result_content = doc_obj.export_to_html()
                elif effective_output_format == "json":
                    result_content = _json(doc_obj.export_to_dict())
                elif effective_output_format == "doctags":
                    result_content = doc_obj.export_to_doctags()
                else:
                    logger.warning(
                        f"Unsupported format '{effective_output_format}' for Docling, using markdown."
                    )
                    result_content = doc_obj.export_to_markdown()
                    effective_output_format = "markdown"

                if save_to_file:
                    fp = (
                        Path(output_path)
                        if output_path
                        else _tmp_path(str(current_input_path), effective_output_format)
                    )
                    fp.parent.mkdir(parents=True, exist_ok=True)
                    img_mode = (
                        _ImageRefModeType.PLACEHOLDER
                        if effective_output_format in ["text", "json"]
                        else _ImageRefModeType.REFERENCED
                    )
                    save_func_map = {
                        "markdown": functools.partial(
                            doc_obj.save_as_markdown, image_mode=img_mode
                        ),
                        "text": functools.partial(doc_obj.save_as_markdown(strict_text=True)),
                        "html": functools.partial(doc_obj.save_as_html, image_mode=img_mode),
                        "json": functools.partial(doc_obj.save_as_json, image_mode=img_mode),
                        "doctags": functools.partial(doc_obj.save_as_doctags),
                    }
                    save_func = save_func_map.get(effective_output_format)
                    if save_func and callable(save_func):
                        with _span(f"docling_save_{effective_output_format}"):
                            save_func(fp)
                        logger.info(f"Saved Docling output ({effective_output_format}) to {fp}")
                        doc_metadata["saved_output_path"] = str(fp)
                    else:
                        fp.write_text(str(result_content), encoding="utf-8")
                        logger.info(f"Saved Docling output (generic text write) to {fp}")
                        doc_metadata["saved_output_path"] = str(fp)

            elif strategy.startswith("fallback_"):
                fallback_type = strategy.split("_", 1)[1]
                logger.info(f"Using fallback strategy for: {fallback_type}")
                fallback_result: Optional[Dict[str, Any]] = None
                if fallback_type == "docx":
                    fallback_result = await _fallback_convert_docx(current_input_path)
                elif fallback_type == "pdf":
                    fallback_result = await _fallback_convert_pdf(current_input_path)
                elif fallback_type == "text":
                    fallback_result = await _fallback_convert_text(current_input_path)
                if fallback_result:
                    raw_text_pages = [fallback_result.get("content", "")]
                    doc_metadata = fallback_result.get("metadata", {})
                    total_doc_pages = doc_metadata.get("num_pages", 1)
                    strategy_used = f"fallback_{fallback_type}"
                else:
                    raise ToolError(
                        "CONVERSION_FAILED",
                        details={"reason": f"Fallback '{fallback_type}' failed."},
                    )

            else:  # Text/OCR strategies
                run_ocr = False
                run_direct = False
                if strategy == "direct_text":
                    run_direct = True
                elif strategy == "ocr":
                    run_ocr = True
                elif strategy == "hybrid_direct_ocr":
                    if not is_pdf:
                        run_ocr = True
                        strategy_used = "ocr"
                        logger.info("Input is image, using 'ocr'.")
                    else:
                        run_direct = True
                extract_start_page = pages_to_process[0] if pages_to_process else 0
                extract_page_count = len(pages_to_process) if pages_to_process else 0

                if run_direct:
                    logger.info(f"Attempting 'direct_text' strategy for {input_name}")
                    try:
                        with _span("direct_text_extraction"):
                            (
                                extracted_pages,
                                has_meaningful_text,
                            ) = await asyncio.to_thread(  # Use helper defined above
                                _ocr_extract_text_from_pdf_direct,
                                current_input_path,
                                start_page=extract_start_page,
                                max_pages=extract_page_count,
                            )
                        total_doc_pages = len(
                            extracted_pages
                        )  # Page count reflects extracted range
                        if strategy == "hybrid_direct_ocr" and not has_meaningful_text:
                            logger.warning("Direct text minimal. Falling back to OCR.")
                            run_ocr = True
                            strategy_used = "ocr"
                        elif not has_meaningful_text and strategy == "direct_text":
                            raise ToolError(
                                "DIRECT_EXTRACTION_FAILED",
                                details={"reason": "No meaningful text found."},
                            )
                        else:
                            raw_text_pages = extracted_pages
                            logger.info(f"Direct text success: {len(raw_text_pages)} pages.")
                    except ToolError as e:
                        if strategy == "hybrid_direct_ocr":
                            logger.warning(f"Direct failed ({e.error_code}). Falling back to OCR.")
                            run_ocr = True
                            strategy_used = "ocr"
                        else:
                            raise e
                    except Exception as e_direct:
                        logger.error(f"Unexpected direct text error: {e_direct}", exc_info=True)
                        if strategy == "hybrid_direct_ocr":
                            logger.warning("Direct failed. Falling back to OCR.")
                            run_ocr = True
                            strategy_used = "ocr"
                        else:
                            raise ToolError(
                                "DIRECT_EXTRACTION_FAILED", details={"error": str(e_direct)}
                            ) from e_direct

                if run_ocr:
                    logger.info(f"Using 'ocr' strategy for {input_name}")
                    strategy_used = "ocr"
                    ocr_lang = ocr_options.get("language", "eng")
                    ocr_dpi = ocr_options.get("dpi", 300)
                    ocr_prep_opts = ocr_options.get("preprocessing")
                    images: List["PILImage.Image"] = []
                    if is_pdf:
                        convert_func = _ocr_convert_pdf_to_images  # Use helper defined above
                        with _span("pdf_to_images"):
                            images = await asyncio.to_thread(
                                convert_func,
                                current_input_path,
                                start_page=extract_start_page,
                                max_pages=extract_page_count,
                                dpi=ocr_dpi,
                            )
                        total_doc_pages = len(images)
                    elif is_image:
                        _ocr_check_dep("Pillow", _PIL_AVAILABLE, "Image loading")
                        if Image is None:
                            raise ToolError(
                                "INTERNAL_ERROR", details={"reason": "PIL.Image is None"}
                            )
                        with _span(f"load_image_{input_name}"):
                            img = Image.open(current_input_path)  # type: ignore
                        images = [img.convert("RGB")]
                        total_doc_pages = 1
                        img.close()  # Close after converting
                    if not images:
                        raise ToolError("OCR_FAILED", details={"reason": "No images for OCR."})

                    processed_pages_text: List[str] = [""] * len(images)

                    async def _process_ocr_page_worker(
                        idx: int, img: "PILImage.Image"
                    ) -> Tuple[int, str]:
                        try:
                            loop = asyncio.get_running_loop()
                            with _span(f"ocr_page_{idx}_preprocess"):
                                prep_img = await loop.run_in_executor(
                                    None, _ocr_preprocess_image, img, ocr_prep_opts
                                )
                            with _span(f"ocr_page_{idx}_tesseract"):
                                text = await loop.run_in_executor(
                                    None,
                                    _ocr_run_tesseract,
                                    prep_img,
                                    ocr_lang,
                                    ocr_options.get("tesseract_config", ""),
                                )  # Use helper defined above
                            if prep_img != img:
                                prep_img.close()  # Close preprocessed image if different
                            return idx, text
                        except Exception as page_err:
                            logger.error(
                                f"OCR page {idx + extract_start_page} error: {page_err}",
                                exc_info=True,
                            )
                            return idx, f"[Page {idx + extract_start_page + 1} OCR Error]"
                        finally:
                            img.close()  # Close the original image passed to worker

                    tasks = [_process_ocr_page_worker(i, img) for i, img in enumerate(images)]
                    page_results = await asyncio.gather(*tasks)
                    for idx, text in page_results:
                        processed_pages_text[idx] = text
                    raw_text_pages = processed_pages_text
                    logger.info(f"OCR extraction successful for {len(raw_text_pages)} pages.")

            # --- Post-Processing (common to all strategies) ---
            if raw_text_pages:
                final_raw_text = "\n\n".join(raw_text_pages)
                if section_filter:
                    # Simple section filtering by keyword
                    filtered_pages = []
                    for page_text in raw_text_pages:
                        if section_filter.lower() in page_text.lower():
                            filtered_pages.append(page_text)
                    final_raw_text = "\n\n".join(filtered_pages)
                    logger.info(f"Applied section filter '{section_filter}', {len(filtered_pages)} pages match.")

                # LLM Enhancement if requested
                if enhance_with_llm and final_raw_text.strip():
                    try:
                        with _span("llm_enhancement"):
                            enhanced_text = await _ocr_enhance_text_chunk(
                                final_raw_text, effective_output_format, ocr_options.get("remove_headers", False)
                            )
                        result_content = enhanced_text
                        logger.info("LLM enhancement completed successfully.")
                    except Exception as e_enhance:
                        logger.warning(f"LLM enhancement failed: {e_enhance}. Using raw text.")
                        result_content = final_raw_text
                else:
                    result_content = final_raw_text

            # Final metadata and quality assessment
            if not doc_metadata:
                doc_metadata = _get_basic_metadata(str(result_content), total_doc_pages)
            
            doc_metadata.update({
                "extraction_strategy_used": strategy_used,
                "output_format": effective_output_format,
                "processing_time_seconds": round(time.time() - t0, 2),
                "enhanced_with_llm": enhance_with_llm,
                "pages_processed": total_doc_pages,
            })

            # Quality assessment if requested
            if ocr_options.get("assess_quality", False) and enhance_with_llm and final_raw_text:
                try:
                    with _span("quality_assessment"):
                        quality_metrics = await _ocr_assess_text_quality(final_raw_text, str(result_content))
                    doc_metadata["quality_assessment"] = quality_metrics
                except Exception as e_quality:
                    logger.warning(f"Quality assessment failed: {e_quality}")

    except ToolError:
        raise
    except Exception as e:
        logger.error(f"Document conversion failed: {e}", exc_info=True)
        raise ToolError("CONVERSION_FAILED", details={"error": str(e)}) from e

    return {
        "content": result_content,
        "metadata": doc_metadata,
        "format": effective_output_format,
        "processing_time": round(time.time() - t0, 2),
        "success": True,
    }


@with_tool_metrics
@with_error_handling
async def chunk_document(
    document: str,
    chunk_size: int = 1000,
    chunk_method: str = "paragraph",
    chunk_overlap: int = 0,
    chunk_strategy: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Split document text into chunks using various strategies.
    
    Args:
        document: Text content to chunk.
        chunk_size: Target maximum size of each chunk (meaning depends on method: tokens or characters).
        chunk_method: Chunking method ('token', 'character', 'section', 'paragraph').
        chunk_overlap: Number of tokens/characters to overlap between chunks.
        chunk_strategy: Alias for chunk_method (for backward compatibility).
    
    Returns:
        Dictionary containing list of chunked text strings.
    """
    if chunk_strategy:
        chunk_method = chunk_strategy
    
    chunks = []
    
    if chunk_method == "paragraph":
        # Split by paragraphs
        paragraphs = document.split('\n\n')
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) + 2 <= chunk_size:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = para
        
        if current_chunk:
            chunks.append(current_chunk)
    
    elif chunk_method == "character":
        # Split by character count
        for i in range(0, len(document), chunk_size - chunk_overlap):
            chunk = document[i:i + chunk_size]
            chunks.append(chunk)
    
    elif chunk_method == "token":
        # Simple token-based splitting (rough approximation: 4 chars per token)
        approx_tokens = len(document) // 4
        token_chunk_size = chunk_size * 4
        for i in range(0, len(document), token_chunk_size - (chunk_overlap * 4)):
            chunk = document[i:i + token_chunk_size]
            chunks.append(chunk)
    
    else:  # Default to sentence-based
        sentences = document.split('. ')
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) + 2 <= chunk_size:
                if current_chunk:
                    current_chunk += ". " + sentence
                else:
                    current_chunk = sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk)
    
    return {
        "chunks": chunks,
        "success": True,
        "chunk_count": len(chunks),
        "method": chunk_method,
        "chunk_size": chunk_size,
        "overlap": chunk_overlap
    }


@with_tool_metrics
@with_error_handling
async def summarize_document(
    document: str,
    max_length: int = 150,
    focus: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Generates a concise summary of the document text using an LLM.
    
    Args:
        document: Text content to summarize.
        max_length: Maximum length of the summary in words.
        focus: Optional focus area for the summary.
    
    Returns:
        Dictionary containing the summary and metadata.
    """
    if not document.strip():
        return {
            "summary": "",
            "word_count": 0,
            "success": True,
            "focus": focus
        }
    
    # Build the prompt
    prompt = f"Please provide a concise summary of the following document in approximately {max_length} words"
    if focus:
        prompt += f", focusing on {focus}"
    prompt += ":\n\n" + document
    
    try:
        # Use the standalone LLM call function
        summary = await _standalone_llm_call(
            prompt=prompt,
            provider=Provider.OPENAI.value,
            temperature=0.3,
            max_tokens=max_length * 2  # Rough approximation
        )
        
        word_count = len(summary.split())
        
        return {
            "summary": summary.strip(),
            "word_count": word_count,
            "original_length": len(document),
            "compression_ratio": round(len(document) / len(summary), 2) if summary else 0,
            "focus": focus,
            "success": True
        }
    
    except Exception as e:
        logger.error(f"Failed to generate summary: {e}")
        # Fallback to simple truncation
        words = document.split()[:max_length]
        fallback_summary = " ".join(words) + "..." if len(words) == max_length else " ".join(words)
        
        return {
            "summary": fallback_summary,
            "word_count": len(words),
            "original_length": len(document),
            "compression_ratio": round(len(document) / len(fallback_summary), 2),
            "focus": focus,
            "success": True,
            "fallback": True,
            "error": str(e)
        }
